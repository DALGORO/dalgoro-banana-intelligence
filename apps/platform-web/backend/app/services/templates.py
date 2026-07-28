from typing import List, Tuple, Optional, Dict, Any
from sqlalchemy.orm import Session
from fastapi import HTTPException
from datetime import datetime
import json
import io, zipfile
from jinja2 import Template

from collections import defaultdict
from statistics import mean
from sqlalchemy import text

import re
_VERSION_RE = re.compile(r"-V(\d{3})$")

import logging
log = logging.getLogger(__name__)

# NUEVO (seguro, con fallback)
from pathlib import Path
try:
    from docxtpl import DocxTemplate
except Exception:
    DocxTemplate = None

from app.models.document_template import DocumentTemplate
from app.models.document import Document
from app.services.storage import save_bytes
from app.services.sst_rules import normalize_activity
from app.models.iperc_item import IPERCItem
from app.services.iperc_derivatives import _row_to_dict
from app.services.known_fields import collect_known_fields


PSICO_DIMS_ORDER = [
  "CARGA Y RITMO DE TRABAJO",
  "DESARROLLO DE COMPETENCIAS",
  "LIDERAZGO",
  "MARGEN DE ACCIÓN Y CONTROL",
  "ORGANIZACIÓN DEL TRABAJO",
  "RECUPERACIÓN",
  "SOPORTE Y APOYO",
  "OTROS PUNTOS IMPORTANTES",
]

def _class_from_avg(x, low_lt=1.5, med_lt=2.5):
  if x < low_lt: return "ALTO"
  if x < med_lt: return "MEDIO"
  return "BAJO"

def fetch_psico_responses(db: Session, company_id: int):
    rows = db.execute(
        text("""
            SELECT worker_index, fields_json AS fields, respuestas_json AS respuestas
            FROM psico_responses
            WHERE company_id = :cid
            ORDER BY worker_index
        """),
        {"cid": company_id},
    ).mappings().all()
    return [dict(r) for r in rows] if rows else []


def build_psico_context(db, company_id: int, data: dict, template_defaults: dict):
  items = fetch_psico_responses(db, company_id)
  if not items and isinstance(data.get("psico_respuestas"), list):
    items = [{
      "worker_index": 1,
      "fields": {k: data.get(k) for k in (
        "fecha","provincia","ciudad","area_trabajo","nivel_instruccion","antiguedad","edad","etnia","sexo"
      ) if data.get(k) is not None},
      "respuestas": data["psico_respuestas"],
    }]

  th = (template_defaults or {}).get("psico_thresholds", {})
  low_lt = float(th.get("low_lt", 1.5))
  med_lt = float(th.get("med_lt", 2.5))

  per_worker = []
  for it in items:
    dims = defaultdict(list)
    for r in (it.get("respuestas") or []):
      try:
        dims[str(r["dimension"]).strip().upper()].append(int(r["puntuacion"]))
      except Exception:
        pass
    rows = []
    for d in PSICO_DIMS_ORDER:
      vals = dims.get(d, [])
      s = sum(vals); n = len(vals)
      prom = (s / n) if n else 0.0
      rows.append({
        "dimension": d,
        "n": n,
        "suma": s,
        "promedio": round(prom, 2),
        "riesgo": _class_from_avg(prom, low_lt, med_lt),
      })
    per_worker.append({
      "worker_index": it.get("worker_index"),
      "rows": rows,
      "total": sum(r["suma"] for r in rows),
      "avg_global": round(mean([r["promedio"] for r in rows if r["n"]]), 2) if any(r["n"] for r in rows) else 0.0,
    })

  dims_all = defaultdict(list)
  for pw in per_worker:
    for r in pw["rows"]:
      if r["n"]:
        dims_all[r["dimension"]].append(r["promedio"])

  diag_table = []
  for d in PSICO_DIMS_ORDER:
    arr = dims_all.get(d, [])
    prom = mean(arr) if arr else 0.0
    diag_table.append({
      "dimension": d.title(),
      "promedio": round(prom, 2),
      "riesgo": _class_from_avg(prom, low_lt, med_lt)
    })

  altos = [r["dimension"] for r in diag_table if r["riesgo"] == "ALTO"]
  medios = [r["dimension"] for r in diag_table if r["riesgo"] == "MEDIO"]
  resumen_txt = []
  if altos: resumen_txt.append(f"Riesgo ALTO en: {', '.join(altos)}.")
  if medios: resumen_txt.append(f"Riesgo MEDIO en: {', '.join(medios)}.")
  if not (altos or medios): resumen_txt.append("Todas las dimensiones en riesgo BAJO.")

  acciones_map = (template_defaults or {}).get("psico_acciones", {}) or {}
  plan_table = []
  for d in altos + medios:
    for acc in (acciones_map.get(d) or acciones_map.get(d.upper()) or acciones_map.get(d.title()) or []):
        plan_table.append({
        "dimension": d.title(),
        "accion": acc,
        "responsable": "",
        "plazo": "",
        "recursos": ""
      })

  return {
    "psico_diag_resumen_txt": " ".join(resumen_txt),
    "psico_diag_table": diag_table,
    "psico_riesgos_txt": "Se recomienda priorizar acciones en dimensiones con riesgo MEDIO/ALTO.",
    "psico_plan_table": plan_table,
    "psico_por_trabajador": per_worker,
    "psico_n_encuestas": len(per_worker),
  }


# --- helpers PSICO (nivel módulo) ---
def _psico_classify(avg: float, thr: dict) -> str:
    """
    Escala 1..4 (4 = mejor, menos riesgo):
      - Bajo  si promedio >= low_ge
      - Medio si promedio >= med_ge
      - Alto  en caso contrario
    """
    low_ge = float((thr or {}).get("low_ge", 3.25))
    med_ge = float((thr or {}).get("med_ge", 2.25))
    return "BAJO" if avg >= low_ge else ("MEDIO" if avg >= med_ge else "ALTO")



def _compute_psico(context: dict) -> dict:
    import json as _json_local  # evita shadowing accidental
    respuestas = context.get("psico_respuestas") or []
    if isinstance(respuestas, str):
        s = respuestas.strip()
        try:
            j = _json_local.loads(s)
            if isinstance(j, list):
                respuestas = j
        except Exception:
            respuestas = []

    if not isinstance(respuestas, list) or not respuestas:
        return {}

    thr = context.get("psico_thresholds") or {"class": "promedio", "low_lt": 1.5, "med_lt": 2.5}

    # 🔧 Normaliza keys: si vienen low_lt/med_lt (escala 1..4, 4=mejor),
    # mapea a low_ge/med_ge que usa _psico_classify (Bajo/Medio/Alto)
    if ("low_ge" not in thr or "med_ge" not in thr) and ("low_lt" in thr or "med_lt" in thr):
        thr = {
            "low_ge": float(thr.get("med_lt", 2.5)),  # Bajo si promedio >= med_lt
            "med_ge": float(thr.get("low_lt", 1.5)),  # Medio si promedio >= low_lt
        }

    acciones_map = context.get("psico_acciones") or {}


    def _to_num(x):
        if isinstance(x, (int, float)):
            return float(x)
        try:
            return float(str(x).strip().replace(",", "."))
        except Exception:
            return None

    buckets, gsum, gn = {}, 0.0, 0
    for r in respuestas:
        dim = str(r.get("dimension", "")).strip()
        if not dim:
            continue
        tipo = str(r.get("tipo", "directa")).strip().lower()
        p = _to_num(r.get("puntuacion"))
        if p is None:
            continue
        if tipo.startswith("inver"):
            p = 5 - p
        b = buckets.setdefault(dim, {"n": 0, "sum": 0.0})
        b["n"] += 1
        b["sum"] += p
        gsum += p
        gn += 1

    dims = []
    for dim, ag in buckets.items():
        n = ag["n"] or 1
        avg = ag["sum"] / n
        dims.append({"factor": dim, "nivel": _psico_classify(avg, thr), "promedio": round(avg, 2), "n": ag["n"]})
    dims.sort(key=lambda x: x["factor"].lower())

    global_avg = round((gsum / gn) if gn else 0.0, 2)
    global_lvl = _psico_classify(global_avg, thr)

    plan = []
    for d in dims:
        medidas = (acciones_map.get(d["factor"]) or acciones_map.get(d["factor"].upper()) or [])
        for m in medidas:
            plan.append({
                "medida": m, "factor": d["factor"],
                "responsable": context.get("responsable_sst") or "SST",
                "recursos": "Recursos internos",
                "f_inicio": str(context.get("fecha_emision") or ""), "f_fin": "",
                "evidencia": "Registro/Acta/Capacitación"
            })

    return {
        "psico_dimensiones": dims,
        "psico_nivel_global": global_lvl,
        "psico_promedio_global": global_avg,
        "psico_resumen_txt": "\n".join(
            f"• {d['factor']}: {d['nivel']} (prom. {d['promedio']}, n={d['n']})" for d in dims
        ),
        "psico_plan_sugerido": plan
    }

# === NUEVO: mapeo del cuestionario oficial y marcadores "X" por celda ===

# Rango de ítems por dimensión (numeración 1..58 igual a tu hoja)
_PSICO_DIMENSIONS = {
    "carga": range(1, 5),                # 1-4
    "competencias": range(5, 9),         # 5-8
    "liderazgo": range(9, 15),           # 9-14
    "margen": range(15, 19),             # 15-18
    "organizacion": range(19, 25),       # 19-24
    "recuperacion": range(25, 30),       # 25-29
    "soporte": range(30, 35),            # 30-34
    "otros": range(35, 59),              # 35-58
}

# Rangos de clasificación exactos de tu tabla
_PSICO_THRESHOLDS = {
    "carga":        [(13,16,"Bajo"), (8,12,"Medio"), (4,7,"Alto")],
    "competencias": [(13,16,"Bajo"), (8,12,"Medio"), (4,7,"Alto")],
    "liderazgo":    [(18,24,"Bajo"), (12,17,"Medio"), (6,11,"Alto")],
    "margen":       [(13,16,"Bajo"), (8,12,"Medio"), (4,7,"Alto")],
    "organizacion": [(18,24,"Bajo"), (12,17,"Medio"), (6,11,"Alto")],
    "recuperacion": [(16,20,"Bajo"), (10,15,"Medio"), (5,9,"Alto")],
    "soporte":      [(16,20,"Bajo"), (10,15,"Medio"), (5,9,"Alto")],
    "otros":        [(73,96,"Bajo"), (49,72,"Medio"), (24,48,"Alto")],
    "global":       [(175,232,"Bajo"), (117,174,"Medio"), (58,116,"Alto")],
}

def _classify_score(score: int, rules: List[Tuple[int,int,str]]) -> str:
    for lo, hi, label in rules:
        if lo <= score <= hi:
            return label
    return ""

def _coerce_list_of_answers(val):
    import json as _json_local
    if isinstance(val, list):
        return val
    if isinstance(val, str):
        s = val.strip()
        try:
            j = _json_local.loads(s)
            return j if isinstance(j, list) else []
        except Exception:
            return []
    return []

def _mark_choice(ctx: dict, prefix: str, value: str, options: Tuple[str, ...]):
    v = (value or "").upper().strip()
    for opt in options:
        ctx[f"{prefix}_{opt}"] = "X" if v == opt else ""

def _prepare_psico_marks(ctx: dict) -> dict:
    """
    - Espera ctx['psico_respuestas'] como lista de {nr, valor} (1..58, 1..4)
    - Genera marcadores 'X' por celda: q01_4, q01_3, q01_2, q01_1 (y también q1_*)
    - Calcula sumas por dimensión y resultado global + etiquetas de riesgo.
    - Marca 'X' en Datos Generales.
    """
    # 1) Normalizar respuestas
    respuestas = _coerce_list_of_answers(ctx.get("psico_respuestas"))
    sums = {k: 0 for k in _PSICO_DIMENSIONS.keys()}
    total = 0

    # 2) “X” por celda y sumatorias
    for item in respuestas:
        try:
            nr = int(item.get("nr") or item.get("n") or item.get("id") or 0)
            val = int(item.get("valor") or item.get("value") or item.get("score") or 0)
        except Exception:
            continue
        if not (1 <= nr <= 58 and 1 <= val <= 4):
            continue

        total += val
        # dos formatos de placeholder (con y sin cero a la izquierda)
        for v in (1, 2, 3, 4):
            ctx[f"q{nr:02d}_{v}"] = "X" if v == val else ""
            ctx[f"q{nr}_{v}"]    = "X" if v == val else ""

        # Acumular en su dimensión
        for dim, rng in _PSICO_DIMENSIONS.items():
            if nr in rng:
                sums[dim] += val
                break

    # 3) Clasificaciones
    for dim, s in sums.items():
        ctx[f"sum_{dim}"] = s
        ctx[f"riesgo_{dim}"] = _classify_score(s, _PSICO_THRESHOLDS[dim]).upper()

    ctx["resultado_global_puntaje"] = total
    ctx["resultado_global_riesgo"] = _classify_score(total, _PSICO_THRESHOLDS["global"]).upper()

    # 4) Datos generales: texto y “X”
    ctx["dg_fecha"]     = str(ctx.get("fecha") or ctx.get("fecha_aplicacion") or ctx.get("fecha_emision") or "")
    ctx["dg_provincia"] = str(ctx.get("provincia") or ctx.get("provincia_aplicacion") or "")
    ctx["dg_ciudad"]    = str(ctx.get("ciudad") or ctx.get("ciudad_aplicacion") or "")
    _mark_choice(ctx, "D_area", ctx.get("area_trabajo"), ("ADM","OP"))
    _mark_choice(ctx, "E_nivel", ctx.get("nivel_instruccion"),
                 ("NINGUNO","BASICA","MEDIA","BACHILLER","TEC","TERCER","CUARTO","OTRO"))
    _mark_choice(ctx, "F_ant", ctx.get("antiguedad"),
                 ("A_0_2","A_3_10","A_11_20","A_21_MAS"))
    _mark_choice(ctx, "G_edad", ctx.get("edad"),
                 ("E_16_24","E_25_34","E_35_43","E_44_52","E_53_MAS"))
    _mark_choice(ctx, "H_etnia", ctx.get("etnia"),
                 ("INDIGENA","AFRO","MESTIZO","BLANCO","MONTUBIO","OTRO"))
    _mark_choice(ctx, "I_sexo", ctx.get("sexo"),
                 ("MUJER","HOMBRE"))

    # 5) Observaciones (ítem 59)
    ctx["obs_59"] = str(ctx.get("observaciones") or "")

    return ctx


def get_template(
    db: Session,
    activity: str,
    code: str,
    risk_level: Optional[str] = None,
    workers: Optional[int] = None,
) -> DocumentTemplate:
    """
    Busca plantilla por actividad normalizada y code con segmentación opcional:
      1) activity + code + risk_level + rango workers (si existen esas columnas)
      2) activity + code + risk_level
      3) activity + code  (comportamiento actual)
      4) GENERICA / GENERICO (fallback actual)

    NOTA: Todo es retrocompatible. Si no pasas risk_level/workers
    se toman como None y aplica el comportamiento actual.
    """
    act = normalize_activity(activity)
    risk = (risk_level or "").strip().upper()

    # Detectar de forma segura si el modelo tiene estas columnas
    has_risk = hasattr(DocumentTemplate, "risk_level")
    has_wmin = hasattr(DocumentTemplate, "workers_min")
    has_wmax = hasattr(DocumentTemplate, "workers_max")
    
    order_col = getattr(DocumentTemplate, "created_at", None) or getattr(DocumentTemplate, "id")
       
    # 1) activity + code + risk + workers (rango) — solo si existen columnas
    if risk and has_risk and (workers is not None) and has_wmin and has_wmax:
        tpl = (
            db.query(DocumentTemplate)
            .filter(
                DocumentTemplate.code == code,
                DocumentTemplate.activity == act,
                DocumentTemplate.risk_level == risk,
                DocumentTemplate.workers_min <= workers,
                DocumentTemplate.workers_max >= workers,
            )
            .order_by(order_col.desc())
            .first()
        )
        if tpl:
            return tpl

    # 2) activity + code + risk — solo si existe columna
    if risk and has_risk:
        tpl = (
            db.query(DocumentTemplate)
            .filter(
                DocumentTemplate.code == code,
                DocumentTemplate.activity == act,
                DocumentTemplate.risk_level == risk,
            )
            .order_by(order_col.desc())
            .first()
        )
        if tpl:
            return tpl

    # 3) activity + code — comportamiento actual
    tpl = (
        db.query(DocumentTemplate)
        .filter(DocumentTemplate.code == code, DocumentTemplate.activity == act)
        .order_by(order_col.desc())
        .first()
    )
    if tpl:
        return tpl

    # 4) GENERICA/GENERICO — comportamiento actual
    tpl = (
        db.query(DocumentTemplate)
        .filter(
            DocumentTemplate.code == code,
            DocumentTemplate.activity.in_(("GENERICA", "GENERICO"))
        )
        .order_by(order_col.desc())
        .first()
    )
    if not tpl:
        raise HTTPException(status_code=404, detail=f"Plantilla no encontrada: {act}/{code}")
    return tpl


def _base_from_template(template_code: str) -> str:
    # Base = prefijo antes del primer guion (RHS-01 -> RHS). Si prefieres usar RHS-01, devuelve template_code.
    return template_code.split("-")[0] if "-" in template_code else template_code

def next_version(db: Session, company_id: int, template_code: str) -> int:
    """
    Siguiente versión para (empresa, plantilla), buscando sufijo -VNNN en el title.
    """
    last = (
        db.query(Document)
        .filter(Document.company_id == company_id, Document.kind == template_code)
        .order_by(Document.created_at.desc())
        .first()
    )
    if not last or not getattr(last, "title", None):
        return 1
    m = _VERSION_RE.search(str(last.title))
    if not m:
        return 1
    try:
        return int(m.group(1)) + 1
    except Exception:
        return 1

def compose_code_versioned(template_code: str, when: datetime, version: int) -> str:
    base = _base_from_template(template_code)  # p.ej. "RHS"
    return f"{base}-{when.strftime('%Y-%m-%d')}-V{version:03d}"


def _docx_template_path(activity: str, code: str) -> Path:
    """
    Ruta esperada del .docx:
      app/static/templates/<ACTIVIDAD_NORMALIZADA>/<CODE>.docx
    """
    # ubicamos .../backend/app/static/templates
    base = Path(__file__).resolve().parents[2] / "app" / "static" / "templates"
    act_dir = normalize_activity(activity)
    return base / act_dir / f"{code}.docx"


# NUEVO
def render_xlsx_template(template_path: str, ctx: dict, render_cfg: dict) -> bytes:
    """
    Renderiza una plantilla XLSX con placeholders Jinja y clona la fila-plantilla
    para cada item en ctx[rows_key]. No afecta al flujo DOCX existente.
    """
    from openpyxl import load_workbook
    from openpyxl.styles import PatternFill
    from openpyxl.cell.cell import MergedCell
    from openpyxl.utils import get_column_letter   # ← NUEVO
    from copy import copy as _copy

    wb = load_workbook(template_path)    
   
    def _has_jinja(v):
        return isinstance(v, str) and "{{" in v and "}}" in v

    def _r(v: str, scope: dict):
        if isinstance(v, str) and "{{" in v and "}}" in v:
            try:
                return Template(v).render(scope)
            except Exception:
                return v
        return v

    xlsx_cfg = (render_cfg or {}).get("xlsx", {})
    rows_key = (render_cfg or {}).get("rows_key") or xlsx_cfg.get("rows_key") or "rows"
    rows = ctx.get(rows_key, []) or []
    sheet_name = xlsx_cfg.get("sheet", "MATRIZ")
    row_marker_idx = int(xlsx_cfg.get("row_marker_row", 2))
    
    formula_by_col = (xlsx_cfg.get("formula_by_column") or {})
        
    
    try:
        ws_tbl = wb[sheet_name]
    except KeyError:
        raise HTTPException(status_code=404, detail=f"XLSX sheet not found: '{sheet_name}' en {template_path}")

  
    # --- AUTO-DETECCIÓN ROBUSTA DE LA FILA PLANTILLA (solo fila de datos) ---
    DATA_KEYS = {"process", "job", "task"}  # puedes añadir "hazard", "event", etc.

    def _row_placeholders(ws, r_idx):
        vals = []
        for c in ws[r_idx]:
            v = c.value
            if isinstance(v, str) and "{{" in v and "}}" in v:
                vals.append(v)
        return vals

    def _row_is_data_template(ws, r_idx):
        # cuenta cuántas llaves de DATA_KEYS están en esta fila
        phs = _row_placeholders(ws, r_idx)
        keys_in_row = set()
        for ph in phs:
            # extrae lo interno de {{ key }}
            m = re.findall(r"{{\s*([a-zA-Z0-9_]+)\s*}}", ph)
            for k in m:
                keys_in_row.add(k.strip())
        return len(DATA_KEYS.intersection(keys_in_row)) >= 2  # exige al menos 2 claves de datos

    # Si la fila configurada no es de datos, busca una que sí lo sea
    if (row_marker_idx <= 0
        or row_marker_idx > ws_tbl.max_row
        or not _row_is_data_template(ws_tbl, row_marker_idx)):
        detected = None
        for ridx in range(1, ws_tbl.max_row + 1):
            if _row_is_data_template(ws_tbl, ridx):
                detected = ridx
                break
        if detected:
            row_marker_idx = detected
        else:
            raise RuntimeError(
                f"No se encontró fila de plantilla de datos en '{sheet_name}'. "
                "Asegura que una fila contenga placeholders como {{ process }}, {{ job }}, {{ task }}."
            )

    # 1) Guardar los placeholders crudos...
    raw_tpl_vals = [c.value for c in ws_tbl[row_marker_idx]]
    tpl_row_cells = [c for c in ws_tbl[row_marker_idx]]
    
    # --- Helpers numéricos y numeración ---
    NUMERIC_KEYS = {
        "nd", "ne", "nc", "np_val", "nr_val",
        "rows_count",  # por si lo imprimes en portada
        # añade aquí cualquier otro campo que quieras forzar como número
    }

    def _looks_number(s: str) -> bool:
        # acepta enteros/decimales con coma o punto
        return bool(re.fullmatch(r"\s*[-+]?\d+(?:[.,]\d+)?\s*", str(s or "")))

    def _to_number(value):
        if value is None:
            return None
        if isinstance(value, (int, float)):
            return value
        s = str(value).strip().replace(",", ".")
        if _looks_number(s):
            try:
                f = float(s)
                # si es entero exacto, devuélvelo como int (openpyxl lo marcará tipo 'n' igual)
                return int(f) if abs(f - int(f)) < 1e-9 else f
            except:
                return value
        return value

    def _write_cell_numeric(dst_cell, key_name: str, rendered_value):
        """
        Escribe 'rendered_value' en dst_cell.
        - Si la llave es numérica o el render tiene pinta de número -> lo guarda como int/float
        - Ajusta number_format simple.
        """
        v = rendered_value

        # 🔹 1) Forzar SI/NO en columnas booleanas del IPERC
        BOOLEAN_KEYS = {
            "requires_work_permit",
            "needs_health_surveillance",
            "needs_env_monitoring",
        }
        v_str = str(v).strip().lower()
        if key_name in BOOLEAN_KEYS and (isinstance(v, bool) or v_str in ("true", "false")):
            dst_cell.value = "SI" if (v is True or v_str == "true") else "NO"
            return

        # 🔹 2) Forzar numérico para columnas que lo requieran (como ya lo hacías)
        if (key_name in NUMERIC_KEYS) or _looks_number(v):
            num = _to_number(v)
            dst_cell.value = num
            try:
                if isinstance(num, int):
                    dst_cell.number_format = "0"
                elif isinstance(num, float):
                    dst_cell.number_format = "0.00"
            except:
                pass
        else:
            dst_cell.value = v
            
    
    # 2) Render global, pero saltando la fila plantilla de MATRIZ
       
    for ws in wb.worksheets:
        for r_index, row in enumerate(ws.iter_rows(), start=1):
            # No tocar la fila plantilla de MATRIZ (se clonará luego)
            if ws.title == sheet_name and r_index == row_marker_idx:
                continue

            for cell in row:
                # 1) No escribir en celdas combinadas no-top-left
                if isinstance(cell, MergedCell):
                    continue

                # 2) Solo renderizar celdas que realmente tengan placeholders {{ }}
                if not _has_jinja(cell.value):
                    continue

                # 3) Render seguro
                try:
                    cell.value = _r(cell.value, ctx)
                except Exception:
                    # Si algo sale mal, deja el valor como está, no rompas el libro
                    pass

    # Limpiar debajo de la fila plantilla
    if ws_tbl.max_row > row_marker_idx:
        ws_tbl.delete_rows(row_marker_idx + 1, ws_tbl.max_row - row_marker_idx)

    def _row_scope(base_ctx, row_dict, idx_zero_based: int):
        # idx humano (1…n)
        n = idx_zero_based + 1
        s = {**base_ctx, **row_dict}
        # numeración disponible para la plantilla (usa {{ row_index }} o {{ row_number }} en la columna #)
        s["row_index"] = n
        s["row_number"] = n
        # aliases útiles
        if "nr" not in s and "nr_val" in row_dict: s["nr"] = row_dict["nr_val"]
        if "np" not in s and "np_val" in row_dict: s["np"] = row_dict["np_val"]
        return s

    def _extract_key_name_from_placeholder(raw_placeholder: str) -> str:
        # de "{{ abc_xyz }}" -> "abc_xyz"; si no hay placeholder, devuelve cadena vacía
        if isinstance(raw_placeholder, str):
            m = re.search(r"{{\s*([a-zA-Z0-9_]+)\s*}}", raw_placeholder)
            if m: return m.group(1)
        return ""

    def _render_and_write_row(target_row_idx: int, scope: dict):
        for col_idx, tpl_cell in enumerate(tpl_row_cells, start=1):
            dst = ws_tbl.cell(row=target_row_idx, column=col_idx)
            raw = raw_tpl_vals[col_idx - 1]
            rendered = _r(raw, scope)
            key_name = _extract_key_name_from_placeholder(raw)

            # 🔹 Normaliza listas a texto legible si el placeholder no tiene |join
            if isinstance(rendered, list):
                rendered = ", ".join(str(x) for x in rendered)

            # 🔹 Mapea booleanos a SI/NO (para S/T/U/Y)
            if isinstance(rendered, bool):
                rendered = "SI" if rendered else "NO"

            _write_cell_numeric(dst, key_name, rendered)
            if tpl_cell.has_style:
                dst.font = _copy(tpl_cell.font)
                dst.border = _copy(tpl_cell.border)
                dst.fill = _copy(tpl_cell.fill)
                dst.number_format = dst.number_format or tpl_cell.number_format
                dst.protection = _copy(tpl_cell.protection)
                dst.alignment = _copy(tpl_cell.alignment)

            # ← NUEVO: si hay fórmula definida para esta columna, la imponemos con la fila real
            col_letter = get_column_letter(col_idx)
            f_tpl = formula_by_col.get(col_letter)
            if f_tpl:
                dst.value = "=" + f_tpl.format(row=target_row_idx)

    if rows:
        # 1) Primera fila de datos: SOBREESCRIBE la fila plantilla
        scope0 = _row_scope(ctx, rows[0], 0)
        _render_and_write_row(row_marker_idx, scope0)

        # 2) Resto de filas: insertar hacia abajo
        insert_at = row_marker_idx + 1
        for i, r in enumerate(rows[1:], start=1):
            ws_tbl.insert_rows(insert_at)
            scope_i = _row_scope(ctx, r, i)
            _render_and_write_row(insert_at, scope_i)
            insert_at += 1
    else:
        # Sin filas: deja la plantilla como está (útil para debug)
        pass

    # 4) Colorear por nivel de riesgo (opcional)
    color_cfg = xlsx_cfg.get("color_cell_by")
    if color_cfg and rows:
        col_name = color_cfg.get("column")          # p.ej., "nr_level"
        color_map = color_cfg.get("map", {})        # {"I":"#..", ...}

        # Busca la columna por placeholder en la fila plantilla
        wanted = "{{ " + str(col_name) + " }}"
        col_idx = None
        for i, raw in enumerate(raw_tpl_vals, start=1):
            if isinstance(raw, str) and wanted in raw:
                col_idx = i
                break

        if col_idx is not None:
            for rr in range(row_marker_idx, ws_tbl.max_row + 1):
        

                lvl = ws_tbl.cell(rr, col_idx).value
                hexcolor = (color_map.get(lvl) or "").replace("#", "")
                if hexcolor:
                    ws_tbl.cell(rr, col_idx).fill = PatternFill(
                        start_color=hexcolor, end_color=hexcolor, fill_type="solid"
                    )

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


def _load_template_render_cfg(tpl: DocumentTemplate) -> dict:
    """
    Lee el sidecar JSON para obtener:
    - render.engine
    - render.rows_key
    - render.xlsx (sheet, row_marker_row, color_cell_by, etc.)
    """
    try:
        base = Path(__file__).resolve().parents[2] / "app" / "static" / "templates"
        candidates = []
        if getattr(tpl, "file_path", None):
            candidates.append((base / tpl.file_path).with_suffix(".json"))
        candidates.append(base / normalize_activity(tpl.activity) / f"{tpl.code}.json")
        try:
            docx_here = _docx_template_path(tpl.activity, tpl.code)
            candidates.append(docx_here.with_suffix(".json"))
        except Exception:
            pass

        for sc in candidates:
            if sc and sc.exists():
                raw = json.loads(sc.read_text(encoding="utf-8")) or {}
                return (raw.get("render") or {})   # ← clave correcta
    except Exception:
        pass
    return {}


def _load_template_defaults(tpl: DocumentTemplate) -> dict:
    try:
        base = Path(__file__).resolve().parents[2] / "app" / "static" / "templates"
        candidates = []

        # Ruta por file_path si existe (misma lógica actual)
        if getattr(tpl, "file_path", None):
            candidates.append((base / tpl.file_path).with_suffix(".json"))

        # Ruta por actividad/código (misma lógica actual)
        candidates.append(base / normalize_activity(tpl.activity) / f"{tpl.code}.json")

        # NUEVO: “al costado” del DOCX realmente usado
        try:
            docx_here = _docx_template_path(tpl.activity, tpl.code)
            candidates.append(docx_here.with_suffix(".json"))
        except Exception:
            pass

        for sc in candidates:
            if sc and sc.exists():
                raw = json.loads(sc.read_text(encoding="utf-8")) or {}
                return (raw.get("defaults") or {})
    except Exception:
        pass
    return {}



# --- BACKFILL DE CONTROLES PARA IPERC-01 (robusto) ---
import json as _json
import re as _re
import unicodedata as _ud

def _load_iperc_catalog(activity: str) -> dict:
    """
    Carga el catálogo JSON por actividad (p.ej., BANANERA.json) desde rutas típicas.
    """
    act = (activity or "").strip().upper() or "BANANERA"
    bases = [
        Path(__file__).resolve().parents[2] / "app" / "static" / "presets" / "iperc",
        Path(__file__).resolve().parents[2] / "presets" / "iperc",
        Path("/mnt/data"),  # útil en tu entorno actual
    ]
    for base in bases:
        p = base / f"{act}.json"
        if p.exists():
            try:
                return _json.loads(p.read_text(encoding="utf-8")) or {}
            except Exception:
                pass
    return {}

def _strip_accents(s: str) -> str:
    s = str(s or "")
    return "".join(ch for ch in _ud.normalize("NFD", s) if _ud.category(ch) != "Mn")

def _normalize_proc_name(s: str) -> str:
    """
    Normaliza cadena de proceso para comparar:
    - Mayúsculas, sin acentos
    - Quita texto entre paréntesis
    - Reemplaza '/', '-' por espacio
    - Colapsa espacios
    """
    s = _strip_accents(str(s or "")).upper()
    s = _re.sub(r"\(.*?\)", "", s)     # quita (SI APLICA), etc.
    s = s.replace("/", " ").replace("-", " ")
    s = _re.sub(r"\s+", " ", s).strip()
    return s

def _normalize_job_name(s: str) -> str:
    # Menos agresivo para job, pero igual sin acentos y en mayúsculas
    return _strip_accents(str(s or "")).upper().strip()

# Sinónimos de procesos: BD → CATÁLOGO
_PROCESS_SYNONYMS = {
    "SIEMBRA TRASPLANTE": "SIEMBRA",
    "DESHOJE DESCHANCLE DESHIJE": "DESHOJE",
    "TRANSPORTE INTERNO TRASLADO": "TRANSPORTE",
    "ACOPIO LAVADO DE RACIMOS": "LAVADO DESINFECCION",
    "CARGA Y DESPACHO": "LOGISTICA",
    "CORTE DESFLORILLADO": "COSECHA",
    "SELECCION CLASIFICACION": "CONTROL DE CALIDAD",             # o "CONTROL DE CALIDAD"; ajustable si lo prefieres
    "BODEGA ALMACEN": "ALMACEN QUIMICOS",
    "BODEGA ALMACEN INSUMOS": "ALMACEN QUIMICOS",
    "RIEGO FERTIRRIEGO": "FERTIRRIEGO",
    "ADMINISTRACION OFICINA": "ADMINISTRACION",
    "FUMIGACION AEREA": "FUMIGACION AEREA",
    "EMPACADO": "EMPAQUE",                             # empacado ≈ empaque
    "MANTENIMIENTO TALLER": "MANTENIMIENTO",          # fallback a MANTENIMIENTO; si job=Soldador, el match flexible encuentra TALLER
}

# Sinónimos de puesto por PROCESO canónico (normalizado)
_JOB_SYNONYMS_BY_PROCESS = {
    "FERTILIZACION": {
        "FERTIRRIEGO": "FERTILIZANTE",
    },
}

def _canon_proc(name: str) -> str:
    n = _normalize_proc_name(name)
    # quita palabras de poca información para mejorar comparación
    n = " ".join(w for w in n.split() if w not in {"DE", "DEL", "LA", "EL", "Y"})
    # aplica sinónimos si existen
    return _PROCESS_SYNONYMS.get(n, n)

def _as_list(x):
    if x is None: return []
    if isinstance(x, list): return x
    if isinstance(x, (tuple, set)): return list(x)
    return [x]

def _join(x):
    return ", ".join(str(i) for i in (x or []))

def _enrich_iperc_rows_with_catalog(activity: str, rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Para cada fila IPERC: si faltan controles, los completa desde el catálogo por (process, job),
    con coincidencia tolerante (normalización, sinónimos y solapamiento de tokens).
    No sobreescribe si ya hay datos. Crea *_txt y controls_planned_txt. Calcula critical_epp si falta.
    """
    cat = _load_iperc_catalog(activity)
    cat_rows = cat.get("rows") or []

    # Índices por combinación normalizada y por proceso normalizado
    idx_full = {}  # (proc_norm, job_norm) -> row
    idx_proc = {}  # proc_norm -> [rows]
    for r in cat_rows:
        pn = _canon_proc(r.get("process", ""))
        jn = _normalize_job_name(r.get("job", ""))
        idx_full[(pn, jn)] = r
        idx_proc.setdefault(pn, []).append(r)

    def _try_match(row: dict) -> dict:
        p_bd = _canon_proc(row.get("process", ""))
        j_bd = _normalize_job_name(row.get("job", ""))

        # aplica sinónimo de puesto según proceso (si existe)  ✅
        j_map = _JOB_SYNONYMS_BY_PROCESS.get(p_bd, {})
        j_bd = j_map.get(j_bd, j_bd)

        # 1) match exacto (proc+job)
        hit = idx_full.get((p_bd, j_bd))
        if hit:
            return hit


        # 2) match por proceso exacto
        lst = idx_proc.get(p_bd)
        if lst:
            # si hay varios, intenta por tokens del job
            tokens = set(t for t in j_bd.split() if len(t) >= 4)
            if tokens:
                for cand in lst:
                    j_c = _normalize_job_name(cand.get("job", ""))
                    if tokens & set(j_c.split()):
                        return cand
            return lst[0]

        # 3) match flexible por solapamiento de tokens de proceso (para nombres “largos” vs cortos)
        #    ejemplo: "ACOPIO LAVADO DE RACIMOS" ~ "LAVADO DESINFECCION"
        tokens_bd = set(p_bd.split())
        best = None
        best_overlap = 0
        for p_cat, rows_cat in idx_proc.items():
            tokens_cat = set(p_cat.split())
            ov = len(tokens_bd & tokens_cat)
            if ov > best_overlap:
                best_overlap, best = ov, rows_cat
        if best:
            # misma heurística de job
            tokens = set(t for t in j_bd.split() if len(t) >= 4)
            if tokens:
                for cand in best:
                    j_c = _normalize_job_name(cand.get("job", ""))
                    if tokens & set(j_c.split()):
                        return cand
            return best[0]

        return {}

    KEYS = [
        "controls_existing_engineering",
        "controls_existing_admin",
        "controls_existing_epp",
        "controls_planned_engineering",
        "controls_planned_admin",
        "controls_planned_epp",
        "critical_epp",
    ]

    for r in rows or []:
        src = _try_match(r)
        # backfill solo si faltan
        for k in KEYS:
            if not _as_list(r.get(k)):
                r[k] = _as_list(src.get(k))

            # versiones *_txt (para plantillas que las usen)
            r[f"{k}_txt"] = _join(r.get(k))

        # Heurística de EPP crítico si aún no hay
        if not _as_list(r.get("critical_epp")):
            epp_exist = _as_list(r.get("controls_existing_epp"))
            epp_plan  = _as_list(r.get("controls_planned_epp"))
            r["critical_epp"] = epp_exist[:1] or epp_plan[:1] or []
            r["critical_epp_txt"] = _join(r["critical_epp"])

        # Texto combinado de planificados (útil para columna R si hay un solo placeholder)
        planned_combo = (
            _as_list(r.get("controls_planned_engineering"))
            + _as_list(r.get("controls_planned_admin"))
            + _as_list(r.get("controls_planned_epp"))
        )
        r["controls_planned_txt"] = _join(planned_combo)
        r["controls_proposed"] = planned_combo
        r["controls_proposed_txt"] = _join(planned_combo)

    return rows


def render_document(tpl: DocumentTemplate, data: Dict[str, Any]) -> Tuple[bytes, str, str]:
    ctx = {"title": tpl.title, "version": tpl.version, **data}

    # Cargar config de render desde el sidecar
    render_cfg = _load_template_render_cfg(tpl)  # dict con keys: engine, rows_key, xlsx
    engine = (getattr(tpl, "render_engine", None) or render_cfg.get("engine"))

    # 1) XLSX si engine=xlsx o si existe el archivo físico .xlsx
    xlsx_path = _xlsx_template_path(tpl.activity, tpl.code)

    if str(engine).lower() == "xlsx":
        if not xlsx_path.exists():
            raise HTTPException(status_code=404, detail=f"XLSX template not found: {xlsx_path}")

        # 🔹 Backfill SOLO para IPERC-01 antes de renderizar
        if str(tpl.code).strip().upper() == "IPERC-01":
            try:
                rows0 = ctx.get("rows", [])
                ctx["rows"] = _enrich_iperc_rows_with_catalog(tpl.activity, rows0)
            except Exception:
                # no romper la exportación si el catálogo no está o algo falla
                pass

        content = render_xlsx_template(
            str(xlsx_path),
            ctx,
            {"xlsx": render_cfg.get("xlsx", {}), "rows_key": render_cfg.get("rows_key", "rows")}
        )
        return content, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "xlsx"

    if xlsx_path.exists():
        # 🔹 Backfill SOLO para IPERC-01 cuando se detecta plantilla XLSX física
        if str(tpl.code).strip().upper() == "IPERC-01":
            try:
                rows0 = ctx.get("rows", [])
                ctx["rows"] = _enrich_iperc_rows_with_catalog(tpl.activity, rows0)
            except Exception:
                pass

        content = render_xlsx_template(
            str(xlsx_path),
            ctx,
            {"xlsx": render_cfg.get("xlsx", {}), "rows_key": render_cfg.get("rows_key", "rows")}
        )
        return content, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "xlsx"


    # 2) DOCX (sin cambios funcionales)
    docx_path = _docx_template_path(tpl.activity, tpl.code)
    if DocxTemplate is not None and docx_path.exists():
        doc = DocxTemplate(str(docx_path))

        # ==== NUEVO: construir tabla del Plan de Acción como subdocumento (sin {% %}) ====
        if str(tpl.code).strip().upper() == "PPRL-01":
            from docx.shared import Inches, Pt, RGBColor  # para anchos opcionales
            # 1) Garantiza pa_rows aun si no fue calculado antes:
            pa_rows = ctx.get("pa_rows")
            if not pa_rows:
                # Fallback: construir desde rows (misma lógica que ya usas para textos)
                def _as_list(x):
                    if x is None: return []
                    if isinstance(x, list): return x
                    if isinstance(x, (tuple, set)): return list(x)
                    return [x]

                def _controls(r, keys):
                    out = []
                    for k in keys: out += _as_list(r.get(k))
                    return out

                pa_rows = []
                for r in ctx.get("rows", []):
                    for c in _controls(r, [
                        "controls_proposed","controls_planned_engineering","controls_planned_admin","controls_planned_epp"
                    ]):
                        pa_rows.append({
                            "control": c,
                            "riesgo": f"{r.get('hazard_group')} – {r.get('hazard')} ({r.get('process')} – {r.get('job')})",
                            "responsable": ctx.get("responsable_sst") or "Delegado/a de SST",
                            "recursos": "Recursos internos SST / presupuesto del área",
                            "f_inicio": str(ctx.get("fecha_emision") or ""),
                            "f_fin": str(ctx.get("periodo_vigencia") or ""),
                            "evidencia": "Registro: procedimiento/capacitación/acta/inspección",
                        })

                if not pa_rows:
                    pa_rows = [{
                        "control": "No observado",
                        "riesgo": "No aplica",
                        "responsable": ctx.get("responsable_sst") or "Delegado/a de SST",
                        "recursos": "Recursos internos SST",
                        "f_inicio": str(ctx.get("fecha_emision") or ""),
                        "f_fin": str(ctx.get("periodo_vigencia") or ""),
                        "evidencia": "Sin registro",
                    }]

            # 2) Crear subdocumento y armar la tabla (con una fila por acción)
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            sub = doc.new_subdoc()
            tbl = sub.add_table(rows=1, cols=7)

            # 1) Intentar aplicar el estilo de tabla definido en el DOCX
            try:
                tbl.style = "PPRL-PlanAccion"  # <— nombre del estilo que creaste en Word
            except Exception:
                tbl.style = "Table Grid"       # fallback

            headers = [
                "Acción/Control", "Riesgo/Referencia IPERC", "Responsable",
                "Recursos", "Fecha inicio", "Fecha fin", "Evidencia/Registro"
            ]
            for i, h in enumerate(headers):
                cell = tbl.cell(0, i)
                cell.text = h

            # 2) Cargar filas
            for a in pa_rows:
                cells = tbl.add_row().cells
                cells[0].text = str(a.get("control", "") or "")
                cells[1].text = str(a.get("riesgo", "") or "")
                cells[2].text = str(a.get("responsable", "") or "")
                cells[3].text = str(a.get("recursos", "") or "")
                cells[4].text = str(a.get("f_inicio", "") or "")
                cells[5].text = str(a.get("f_fin", "") or "")
                cells[6].text = str(a.get("evidencia", "") or "")

            # 3) Fallback de formato programático si el estilo no existe (colores/fuentes)
            def _shade(cell, hex_color):
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), hex_color)  # p. ej. '1F4E79'
                tcPr.append(shd)

            def _set_cell_font(cell, name="Calibri", size_pt=10, bold=False, rgb=None):
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (cell in [tbl.rows[0].cells[4], tbl.rows[0].cells[5]]) else p.alignment
                    for r in p.runs:
                        r.font.name = name
                        r.font.size = Pt(size_pt)
                        r.font.bold = bold
                        if rgb:
                            r.font.color.rgb = rgb

            # Detecta si el estilo personalizado no se aplicó (muy simple)
            style_used = (getattr(tbl.style, "name", "") == "PPRL-PlanAccion")
            if not style_used:
                # Encabezado: fondo azul oscuro y texto blanco, negrita
                for c in tbl.rows[0].cells:
                    _shade(c, "1F4E79")
                    for p in c.paragraphs:
                        if not p.runs:
                            p.add_run(c.text or "")
                            c.text = ""
                        for r in p.runs:
                            r.font.bold = True
                            r.font.size = Pt(10)
                            r.font.name = "Calibri"
                            r.font.color.rgb = RGBColor(255, 255, 255)

                # Cuerpo: fuente Calibri 10, zebra suave
                for idx, row in enumerate(tbl.rows[1:], start=1):
                    for c in row.cells:
                        _set_cell_font(c, name="Calibri", size_pt=10)
                    if idx % 2 == 0:
                        for c in row.cells:
                            _shade(c, "EAF2FF")  # azul muy claro para bandas

            # 4) Anchos de columna (ajústalos a tu gusto)
            try:
                widths = [2.0, 3.0, 2.2, 2.2, 1.6, 1.6, 2.4]  # pulgadas
                for col_idx, w in enumerate(widths):
                    for r in tbl.rows:
                        r.cells[col_idx].width = Inches(w)
            except Exception:
                pass

            # Encabezados de fechas (col 5 y 6 en la tabla, índices 4 y 5):
            for col_idx in (4, 5):
                for p in tbl.rows[0].cells[col_idx].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Celdas de fechas en todas las filas de datos:
            for row in tbl.rows[1:]:
                for col_idx in (4, 5):
                    for p in row.cells[col_idx].paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # === /Alineación ===
            
            ctx["pa_table"] = sub
        
        elif str(tpl.code).strip().upper() == "CAP-01":
            # ---- Subdocumento 4: Malla de capacitación ----
            
            import json as _json

            def _as_rows(x):
                if isinstance(x, list):
                    return x
                if isinstance(x, str):
                    try:
                        y = _json.loads(x)
                        return y if isinstance(y, list) else []
                    except Exception:
                        return []
                return []

            # Asegura que el contexto tenga listas válidas
            ctx["capacitaciones"] = _as_rows(ctx.get("capacitaciones"))
            ctx["programacion_mensual"] = _as_rows(ctx.get("programacion_mensual"))
            
            sub_malla = doc.new_subdoc()
            tbl_m = sub_malla.add_table(rows=1, cols=13)
            try:
                tbl_m.style = "Table Grid"
            except Exception:
                pass
            headers = [
                "Origen","Tema / Contenido","Referencia IPERC","Población","Nº part.",
                "Periodicidad","Duración (h)","Modalidad","Proveedor","Mes",
                "Evidencia","Indicador","Fuente"
            ]
            for i, h in enumerate(headers):
                tbl_m.cell(0, i).text = h

            for it in (ctx.get("capacitaciones") or []):
                row = tbl_m.add_row().cells
                row[0].text  = str(it.get("origen","") or "")
                row[1].text  = str(it.get("tema","") or "")
                row[2].text  = str(it.get("referencia_iperc","") or "")
                row[3].text  = str(it.get("poblacion","") or "")
                row[4].text  = str(it.get("n_participantes","") or "")
                row[5].text  = str(it.get("periodicidad","") or "")
                row[6].text  = str(it.get("duracion_h","") or "")
                row[7].text  = str(it.get("modalidad","") or "")
                row[8].text  = str(it.get("proveedor","") or "")
                row[9].text  = str(it.get("mes","") or "")
                row[10].text = str(it.get("evidencia","") or "")
                row[11].text = str(it.get("indicador","") or "")
                row[12].text = f"{it.get('source_doc','') or ''} / {it.get('source_section','') or ''}"
                
            if len(tbl_m.rows) == 1:
                r = tbl_m.add_row().cells
                r[0].text = "— Sin datos —"

            ctx["cap_malla_table"] = sub_malla

            # ---- Subdocumento 5: Programación mensual ----
            sub_prog = doc.new_subdoc()
            tbl_p = sub_prog.add_table(rows=1, cols=5)
            try:
                tbl_p.style = "Table Grid"
            except Exception:
                pass
            headers_p = ["Mes","Actividades/Temas (con origen)","Áreas/Puestos","Responsable","Observaciones"]
            for i, h in enumerate(headers_p):
                tbl_p.cell(0, i).text = h

            for prg in (ctx.get("programacion_mensual") or []):
                row = tbl_p.add_row().cells
                row[0].text = str(prg.get("mes","") or "")
                row[1].text = str(prg.get("actividades","") or "")
                row[2].text = str(prg.get("areas","") or "")
                row[3].text = str(prg.get("responsable","") or "")
                row[4].text = str(prg.get("observaciones","") or "")
                
            if len(tbl_p.rows) == 1:
                r = tbl_p.add_row().cells
                r[0].text = "— Sin datos —"

            ctx["cap_programacion_table"] = sub_prog

            # ---- Textos sin bucles para base legal / roles / presupuesto / CC ----
            # base legal
            bl = ctx.get("base_legal") or []
            ctx["base_legal_txt"] = "\n".join(f"• {i}" for i in bl) if bl else "• Sin referencias"

            # roles (tres columnas en paralelo, mismo número de líneas)
            roles = ctx.get("roles") or []
            ctx["roles_rol_txt"]  = "\n".join([str(r.get('rol','')) for r in roles]) or "Alta dirección\nDelegado/a de SST\nMédico ocupacional/Servicio de salud\nSupervisores/Jefaturas"
            ctx["roles_resp_txt"] = "\n".join([str(r.get('responsable','')) for r in roles]) or "\n".join([
                ctx.get("representante_legal",""),
                ctx.get("responsable_sst",""),
                ctx.get("medico_ocupacional",""),
                ctx.get("responsables_area","")
            ])
            ctx["roles_func_txt"] = "\n".join([str(r.get('funciones','')) for r in roles]) or "Aprobar plan y asignar recursos.\nCoordinar ejecución, seguimiento e indicadores.\nVigilancia de la salud.\nGarantizar asistencia y aplicación en el puesto."

            # ---- Presupuesto (robustez) ----
            import json as _json, re as _re

            def _norm_pres(p):
                # 1) ya es lista
                if isinstance(p, list):
                    out = []
                    for el in p:
                        if isinstance(el, dict):
                            out.append(el)
                        elif isinstance(el, str) and el.strip():
                            m = _re.search(r"[-\d.,]+", el)
                            val = float((m.group(0) if m else "0").replace(".", "").replace(",", ".")) if m else 0.0
                            out.append({"concepto": el, "cantidad": 1, "costo_unitario": val, "proveedor": ""})
                    return out
                # 2) viene como string JSON
                if isinstance(p, str):
                    s = p.strip()
                    try:
                        j = _json.loads(s)
                        if isinstance(j, list):
                            return _norm_pres(j)
                    except Exception:
                        pass
                    # 3) string de texto / con valor $ -> lo convierto a 1 ítem
                    m = _re.search(r"[-\d.,]+", s)
                    val = float((m.group(0) if m else "0").replace(".", "").replace(",", ".")) if m else 0.0
                    if s or val > 0:
                        return [{"concepto": "Presupuesto general", "cantidad": 1, "costo_unitario": val, "proveedor": ""}]
                return []

            pres = _norm_pres(ctx.get("presupuesto"))
            lines, total = [], 0.0

            def _to_float(v):
                s = str(v or "").strip().replace(",", ".")
                try: return float(s)
                except: return 0.0

            for it in pres:
                cant = _to_float(it.get("cantidad"))
                unit = _to_float(it.get("costo_unitario"))
                sub  = cant * unit
                total += sub
                prov = str(it.get("proveedor","") or "")
                lines.append(f"{it.get('concepto','')} — {cant} x {unit} = {round(sub,2)} ({prov})")

            ctx["presupuesto_txt"] = "\n".join(lines) if lines else "— Sin presupuesto estimado"
            ctx["presupuesto_total_txt"] = f"USD {round(total,2)}"


            # control de cambios (columnas paralelas)
            cc = ctx.get("control_cambios") or [{
                "version": ctx.get("version","1.0"),
                "fecha": ctx.get("fecha_emision",""),
                "descripcion": "Emisión inicial",
                "aprobo": (ctx.get("representante_legal") or "")
            }]
            ctx["cc_version_txt"] = "\n".join(str(i.get("version","")) for i in cc)
            ctx["cc_fecha_txt"]   = "\n".join(str(i.get("fecha","")) for i in cc)
            ctx["cc_desc_txt"]    = "\n".join(str(i.get("descripcion","")) for i in cc)
            ctx["cc_aprobo_txt"]  = "\n".join(str(i.get("aprobo","")) for i in cc)
            
        # --- Normalización robusta de listas para CAP-01 (por si llegan como string JSON) ---
        
        elif str(tpl.code).strip().upper() == "EPP-01":
            # --- Subdoc A: Matriz de EPP por Proceso/Puesto ---
            sub_mat = doc.new_subdoc()
            t = sub_mat.add_table(rows=1, cols=7)
            try:
                t.style = "Table Grid"
            except Exception:
                pass
            headers = [
                "Proceso","Puesto","Peligros/Exposición",
                "EPP mínimo obligatorio","Reposición","Responsable entrega","Evidencia"
            ]
            for i, h in enumerate(headers):
                t.cell(0, i).text = h

            for it in (ctx.get("epp_matrix") or []):
                row = t.add_row().cells
                row[0].text = str(it.get("process","") or "")
                row[1].text = str(it.get("job","") or "")
                row[2].text = str(it.get("hazards_txt","") or "")
                row[3].text = str(it.get("epp_txt","") or "")
                row[4].text = str(it.get("reposicion","") or "")
                row[5].text = str(it.get("responsable_entrega","") or "")
                row[6].text = str(it.get("evidencia","") or "")

            if len(t.rows) == 1:
                r = t.add_row().cells
                r[0].text = "— Sin datos —"

            ctx["epp_matrix_table"] = sub_mat

            # --- Subdoc B: Registros de entrega (formato de firmas) ---
            sub_reg = doc.new_subdoc()
            tr = sub_reg.add_table(rows=1, cols=8)
            try:
                tr.style = "Table Grid"
            except Exception:
                pass
            headers_r = [
                "Nombre","Puesto","EPP entregado","Talla","Cantidad","Fecha",
                "Firma trabajador","Firma responsable"
            ]
            for i, h in enumerate(headers_r):
                tr.cell(0, i).text = h

            for it in (ctx.get("epp_registros") or []):
                r = tr.add_row().cells
                r[0].text = str(it.get("nombre","") or "")
                r[1].text = str(it.get("puesto","") or "")
                r[2].text = str(it.get("epp_entregado","") or "")
                r[3].text = str(it.get("talla","") or "")
                r[4].text = str(it.get("cantidad","") or "")
                r[5].text = str(it.get("fecha","") or "")
                r[6].text = str(it.get("firma_trabajador","") or "")
                r[7].text = str(it.get("firma_responsable","") or "")

            if len(tr.rows) == 1:
                r = tr.add_row().cells
                r[0].text = "— Sin datos —"

            ctx["epp_registros_table"] = sub_reg

            # --- Textos auxiliares (idéntico patrón a CAP-01) ---
            # base legal
            bl = ctx.get("base_legal") or []
            ctx["base_legal_txt"] = "\n".join(f"• {i}" for i in bl) if bl else "• Sin referencias"

            # roles en 3 columnas
            roles = ctx.get("roles") or []
            ctx["roles_rol_txt"]  = "\n".join([str(r.get('rol','')) for r in roles]) or "Alta dirección\nDelegado/a de SST\nMédico ocupacional/Servicio de salud\nSupervisores/Jefaturas"
            ctx["roles_resp_txt"] = "\n".join([str(r.get('responsable','')) for r in roles]) or "\n".join([
                ctx.get("representante_legal",""),
                ctx.get("responsable_sst",""),
                ctx.get("medico_ocupacional",""),
                ctx.get("responsables_area","")
            ])
            ctx["roles_func_txt"] = "\n".join([str(r.get('funciones','')) for r in roles]) or "Aprobar plan y asignar recursos.\nCoordinar ejecución, seguimiento e indicadores.\nVigilancia de la salud.\nGarantizar uso y reposición de EPP."

            # control de cambios
            cc = ctx.get("control_cambios") or [{
                "version": ctx.get("version","1.0"),
                "fecha": ctx.get("fecha_emision",""),
                "descripcion": "Emisión inicial",
                "aprobo": (ctx.get("representante_legal") or "")
            }]
            ctx["cc_version_txt"] = "\n".join(str(i.get("version","")) for i in cc)
            ctx["cc_fecha_txt"]   = "\n".join(str(i.get("fecha","")) for i in cc)
            ctx["cc_desc_txt"]    = "\n".join(str(i.get("descripcion","")) for i in cc)
            ctx["cc_aprobo_txt"]  = "\n".join(str(i.get("aprobo","")) for i in cc)

                    
        elif str(tpl.code).strip().upper() == "EMG-01":
            # --- Subdoc 6: Matriz de escenarios ---
            sub_m = doc.new_subdoc()
            tbl = sub_m.add_table(rows=1, cols=9)
            try: tbl.style = "Table Grid"
            except: pass
            headers = ["Escenario","Clasificación","Referencia IPERC","Controles existentes","Medidas de respuesta","Equipos","Roles","Punto de encuentro","Evidencia"]
            for i,h in enumerate(headers): tbl.cell(0,i).text = h

            for it in (ctx.get("emg_escenarios") or []):
                row = tbl.add_row().cells
                row[0].text = str(it.get("escenario",""))
                row[1].text = str(it.get("clasificacion",""))
                row[2].text = str(it.get("referencia_iperc",""))
                row[3].text = str(it.get("controles_existentes",""))
                row[4].text = "Activar alarma, evacuar, cortar energías, atender heridos, contener y comunicar"
                row[5].text = str(it.get("equipos_requeridos",""))
                row[6].text = str(it.get("roles",""))
                row[7].text = str(it.get("punto_encuentro",""))
                row[8].text = str(it.get("evidencia",""))
            if len(tbl.rows) == 1:
                tbl.add_row().cells[0].text = "— Sin datos —"
            ctx["emg_matriz_table"] = sub_m

            # --- Subdoc 7: Programación de simulacros ---
            sub_s = doc.new_subdoc()
            t2 = sub_s.add_table(rows=1, cols=5)
            try: t2.style = "Table Grid"
            except: pass
            for i,h in enumerate(["Mes","Escenario","Objetivo","Responsable","Observaciones"]): t2.cell(0,i).text = h
            for it in (ctx.get("emg_simulacros") or []):
                r = t2.add_row().cells
                r[0].text = str(it.get("mes",""))
                r[1].text = str(it.get("escenario",""))
                r[2].text = str(it.get("objetivo",""))
                r[3].text = str(it.get("responsable",""))
                r[4].text = str(it.get("observaciones",""))
            if len(t2.rows) == 1:
                t2.add_row().cells[0].text = "— Sin datos —"
            ctx["emg_simulacros_table"] = sub_s

            # --- Subdoc 8: Contactos externos ---
            sub_c = doc.new_subdoc()
            t3 = sub_c.add_table(rows=1, cols=4)
            try: t3.style = "Table Grid"
            except: pass
            for i,h in enumerate(["Entidad","Contacto","Teléfono","Observaciones"]): t3.cell(0,i).text = h
            for it in (ctx.get("emg_contactos") or []):
                r = t3.add_row().cells
                r[0].text = str(it.get("entidad",""))
                r[1].text = str(it.get("contacto",""))
                r[2].text = str(it.get("telefono",""))
                r[3].text = str(it.get("obs",""))
            if len(t3.rows) == 1:
                t3.add_row().cells[0].text = "— Sin datos —"
            ctx["emg_contactos_table"] = sub_c

            # --- Control de cambios por defecto (igual que CAP-01) ---
            cc = ctx.get("control_cambios") or [{
                "version": ctx.get("version","1.0"),
                "fecha": ctx.get("fecha_emision",""),
                "descripcion": "Emisión inicial",
                "aprobo": (ctx.get("representante_legal") or "")
            }]
            ctx["cc_version_txt"] = "\n".join(str(i.get("version","")) for i in cc)
            ctx["cc_fecha_txt"]   = "\n".join(str(i.get("fecha","")) for i in cc)
            ctx["cc_desc_txt"]    = "\n".join(str(i.get("descripcion","")) for i in cc)
            ctx["cc_aprobo_txt"]  = "\n".join(str(i.get("aprobo","")) for i in cc)

        
        elif str(tpl.code).strip().upper() == "PSICO-01":
            # ===== Equipo de trabajo =====
            sub_eq = doc.new_subdoc()
            t_eq = sub_eq.add_table(rows=1, cols=4)
            try: t_eq.style = "Table Grid"
            except: pass
            for i, h in enumerate(["Rol","Nombre","Responsabilidades","Contacto"]):
                t_eq.cell(0, i).text = h
            for it in (ctx.get("psico_equipo") or []):
                r = t_eq.add_row().cells
                r[0].text = str(it.get("rol",""))
                r[1].text = str(it.get("nombre",""))
                r[2].text = str(it.get("responsabilidades",""))
                r[3].text = str(it.get("contacto",""))
            if len(t_eq.rows) == 1:
                t_eq.add_row().cells[0].text = "— Sin datos —"
            ctx["equipo_table"] = sub_eq

            # ===== Cronograma (anual) =====
            sub_cr = doc.new_subdoc()
            t_cr = sub_cr.add_table(rows=1, cols=5)
            try: t_cr.style = "Table Grid"
            except: pass
            for i, h in enumerate(["Actividad","Mes","Responsable","Evidencia","Obs."]):
                t_cr.cell(0, i).text = h
            for it in (ctx.get("psico_cronograma") or []):
                r = t_cr.add_row().cells
                r[0].text = str(it.get("actividad",""))
                r[1].text = str(it.get("mes",""))
                r[2].text = str(it.get("responsable",""))
                r[3].text = str(it.get("evidencia",""))
                r[4].text = str(it.get("obs",""))
            if len(t_cr.rows) == 1:
                t_cr.add_row().cells[0].text = "— Sin datos —"
            ctx["cronograma_table"] = sub_cr

            # ===== Resultados (resumen de factores) =====
            sub_rs = doc.new_subdoc()
            t_rs = sub_rs.add_table(rows=1, cols=4)
            try: t_rs.style = "Table Grid"
            except: pass
            for i, h in enumerate(["Factor","Nivel","Hallazgos clave","Población"]):
                t_rs.cell(0, i).text = h
            for it in (ctx.get("psico_resultados") or []):
                r = t_rs.add_row().cells
                r[0].text = str(it.get("factor",""))
                r[1].text = str(it.get("nivel",""))
                r[2].text = str(it.get("hallazgos",""))
                r[3].text = str(it.get("poblacion",""))
            if len(t_rs.rows) == 1:
                t_rs.add_row().cells[0].text = "— Sin datos —"
            ctx["resultados_table"] = sub_rs

            # ===== Plan de intervención =====
            sub_pl = doc.new_subdoc()
            t_pl = sub_pl.add_table(rows=1, cols=7)
            try: t_pl.style = "Table Grid"
            except: pass
            hdrs = ["Medida","Factor ref.","Responsable","Recursos","Inicio","Fin","Evidencia"]
            for i, h in enumerate(hdrs): t_pl.cell(0, i).text = h
            for it in (ctx.get("psico_plan") or []):
                r = t_pl.add_row().cells
                r[0].text = str(it.get("medida",""))
                r[1].text = str(it.get("factor",""))
                r[2].text = str(it.get("responsable",""))
                r[3].text = str(it.get("recursos",""))
                r[4].text = str(it.get("f_inicio",""))
                r[5].text = str(it.get("f_fin",""))
                r[6].text = str(it.get("evidencia",""))
            if len(t_pl.rows) == 1:
                t_pl.add_row().cells[0].text = "— Sin datos —"
            ctx["psico_plan_table"] = sub_pl

            # ===== Indicadores (seguimiento) =====
            sub_ind = doc.new_subdoc()
            t_in = sub_ind.add_table(rows=1, cols=5)
            try: t_in.style = "Table Grid"
            except: pass
            for i, h in enumerate(["Indicador","Fórmula","Meta","Fuente","Frecuencia"]):
                t_in.cell(0, i).text = h
            for it in (ctx.get("psico_indicadores") or []):
                r = t_in.add_row().cells
                r[0].text = str(it.get("indicador",""))
                r[1].text = str(it.get("formula",""))
                r[2].text = str(it.get("meta",""))
                r[3].text = str(it.get("fuente",""))
                r[4].text = str(it.get("frecuencia",""))
            if len(t_in.rows) == 1:
                t_in.add_row().cells[0].text = "— Sin datos —"
            ctx["indicadores_table"] = sub_ind

        
        elif str(tpl.code).strip().upper() == "AGRO-PLAG-01":
            # --- Subdoc A: Inventario de plaguicidas ---
            sub_inv = doc.new_subdoc()
            tbl = sub_inv.add_table(rows=1, cols=8)
            try: tbl.style = "Table Grid"
            except: pass
            headers = ["Producto","Ingrediente activo","Formulación","Registro","Categoría tox.","HDS","Proveedor","Venc."]
            for i,h in enumerate(headers): tbl.cell(0,i).text = h

            _inv_raw = ctx.get("inventario_plaguicidas")
            if isinstance(_inv_raw, str):
                s = _inv_raw.strip()
                if s.startswith("[") and s.endswith("]"):
                    try:
                        ctx["inventario_plaguicidas"] = json.loads(s)
                    except Exception:
                        pass
            
            for it in (ctx.get("inventario_plaguicidas") or []):
                r = tbl.add_row().cells
                r[0].text = str(it.get("producto","") or "")
                r[1].text = str(it.get("ingrediente_activo","") or "")
                r[2].text = str(it.get("formulacion","") or "")
                r[3].text = str(it.get("registro","") or "")
                r[4].text = str(it.get("categoria_tox","") or "")
                r[5].text = str(it.get("hds","") or "")
                r[6].text = str(it.get("proveedor","") or "")
                r[7].text = str(it.get("vencimiento","") or "")

            if len(tbl.rows) == 1:
                tbl.add_row().cells[0].text = "— Sin datos —"
            ctx["inventario_table"] = sub_inv

            # --- Subdoc B: Programación de tratamientos ---
            sub_prog = doc.new_subdoc()
            t2 = sub_prog.add_table(rows=1, cols=6)
            try: t2.style = "Table Grid"
            except: pass
            hdrs = ["Cultivo/Área","Objetivo","Producto","Dosis","Fecha","Responsable"]
            for i,h in enumerate(hdrs): t2.cell(0,i).text = h

            _prog_raw = ctx.get("programacion_tratamientos")
            if isinstance(_prog_raw, str):
                s = _prog_raw.strip()
                if s.startswith("[") and s.endswith("]"):
                    try:
                        ctx["programacion_tratamientos"] = json.loads(s)
                    except Exception:
                        pass
            
            for it in (ctx.get("programacion_tratamientos") or []):
                r = t2.add_row().cells
                r[0].text = str(it.get("cultivo_area","") or "")
                r[1].text = str(it.get("objetivo","") or "")
                r[2].text = str(it.get("producto","") or "")
                r[3].text = str(it.get("dosis","") or "")
                r[4].text = str(it.get("fecha","") or "")
                r[5].text = str(it.get("responsable","") or "")

            if len(t2.rows) == 1:
                t2.add_row().cells[0].text = "— Sin datos —"
            ctx["programacion_table"] = sub_prog

            def _as_list_str(val):
                if isinstance(val, str):
                    return [x.strip() for x in val.splitlines() if x.strip()]
                if isinstance(val, list):
                    return [str(x).strip() for x in val if str(x).strip()]
                return []
            
            # Textos “*_txt” si llegaron listas (compat con otras plantillas)
            def _bullets(lst):
                items = _as_list_str(lst)
                return "\n".join(f"• {x}" for x in items) if items else "• Sin datos"


            if not ctx.get("base_legal_txt"): ctx["base_legal_txt"] = _bullets(ctx.get("base_legal"))
            if not ctx.get("almacenamiento_lineamientos_txt"): ctx["almacenamiento_lineamientos_txt"] = _bullets(ctx.get("almacenamiento_lineamientos"))
            if not ctx.get("procedimientos_escritos_txt"): ctx["procedimientos_escritos_txt"] = _bullets(ctx.get("procedimientos_escritos"))
            if not ctx.get("epp_obligatorio_txt"): ctx["epp_obligatorio_txt"] = _bullets(ctx.get("epp_obligatorio"))
            if not ctx.get("hds_enlaces_txt"): ctx["hds_enlaces_txt"] = _bullets(ctx.get("hds_links"))
            if not ctx.get("capacitacion_temas_txt"): ctx["capacitacion_temas_txt"] = _bullets(ctx.get("capacitacion_temas"))
            if not ctx.get("vigilancia_items_txt"): ctx["vigilancia_items_txt"] = _bullets(ctx.get("vigilancia_items"))
            if not ctx.get("indicadores_txt"): ctx["indicadores_txt"] = _bullets(ctx.get("indicadores"))
            
            cc = ctx.get("control_cambios") or [{
                "version": ctx.get("version", "1.0"),
                "fecha": str(ctx.get("fecha_emision") or ""),
                "descripcion": "Emisión inicial",
                "aprobo": (ctx.get("representante_legal") or "")
            }]
            ctx["cc_version_txt"] = "\n".join(str(i.get("version","")) for i in cc)
            ctx["cc_fecha_txt"]   = "\n".join(str(i.get("fecha","")) for i in cc)
            ctx["cc_desc_txt"]    = "\n".join(str(i.get("descripcion","")) for i in cc)
            ctx["cc_aprobo_txt"]  = "\n".join(str(i.get("aprobo","")) for i in cc)
        

        doc.render(ctx)
        out = io.BytesIO()
        doc.save(out)
        content = out.getvalue()
        return content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"


    # 3) Fallback JSON
    content = json.dumps({"title": tpl.title, "version": tpl.version, "data": data}, ensure_ascii=False, indent=2).encode("utf-8")
    return content, "application/json", "json"


def _xlsx_template_path(activity: str, code: str) -> Path:
    base = Path(__file__).resolve().parents[2] / "app" / "static" / "templates"
    act_dir = normalize_activity(activity)
    return base / act_dir / f"{code}.xlsx"


def generate_document(db, company, template_code, activity, data, risk_level, workers, extra_meta: Optional[dict] = None, user_id: int = 0):
    extra_meta = extra_meta or {}
    code_up = str(template_code).strip().upper()

    # 👇 Normalizar company_id (soporta ORM, dict o id directo)
    company_id = (
        getattr(company, "id", None)
        or (company.get("id") if isinstance(company, dict) else None)
        or (int(company) if isinstance(company, (int, str)) and str(company).isdigit() else None)
    )
    if company_id is None:
        raise HTTPException(status_code=400, detail="company_id not provided")

    # ---- contexto para DOCX / JSON ----
    legal_txt = (extra_meta.get("legal") or "").strip() or None
    context = {
        "company": getattr(company, "name", getattr(company, "razon_social", "")),
        "actividad": activity,
        "riesgo": risk_level,
        "trabajadores": workers,
        **data,
    }
    # Alias seguro de ruc desde la empresa (si no viene en el form)
    context.setdefault("ruc", getattr(company, "ruc", ""))

    # 👇 Cargar plantilla y defaults ANTES de cualquier rama específica
    tpl = get_template(db, activity=activity, code=template_code, risk_level=risk_level, workers=workers)
    try:
        defaults = _load_template_defaults(tpl)
        for k, v in (defaults or {}).items():
            cur = context.get(k)
            if cur is None or (isinstance(cur, str) and not cur.strip()) or cur in ([], {}):
                context[k] = v
    except Exception:
        pass

    # 👇 Solo PSICO: anexar contexto calculado usando company_id y defaults ya cargados
    if str(template_code).upper() == "PSICO-01":
        extra = build_psico_context(db, company_id, data or {}, defaults or {})
        context.update(extra)

    
    # Asegurar al menos 1 HDS antes de la validación (por si el sidecar viene sin defaults)
    if str(template_code).strip().upper() == "AGRO-PLAG-01":
        h = context.get("hds_links")
        has = (isinstance(h, list) and any(str(x).strip() for x in h)) or (isinstance(h, str) and h.strip())
        if not has:
            context["hds_links"] = ["HDS física/digital"]
    
    # --- AUTOGENERACIÓN SOLO PARA AGRO-PLAG-01 en actividad BANANERA ---
    if str(template_code).strip().upper() == "AGRO-PLAG-01" and str(activity or "").strip().upper() == "BANANERA":
        import json as _json, re as _re

        # 1) Normalizadores/validadores mínimos
        def _as_list_dicts(x):
            # si viene string con JSON (textarea), intenta parsear
            if isinstance(x, str):
                s = x.strip()
                if s.startswith("[") and s.endswith("]"):
                    try:
                        j = _json.loads(s)
                        if isinstance(j, list):
                            return [d for d in j if isinstance(d, dict)]
                    except Exception:
                        pass
                return []
            if isinstance(x, list):
                return [d for d in x if isinstance(d, dict)]
            return []

        def _has_text(v): 
            return bool(str(v or "").strip())

        def _valid_inv(items):
            # al menos 1 fila con todos los campos obligatorios poblados (no vacíos)
            req = ("producto","ingrediente_activo","formulacion","registro","categoria_tox","hds")
            for it in items:
                if all(_has_text(it.get(k)) for k in req):
                    return True
            return False

        def _valid_prog(items):
            # al menos 1 fila con campos clave (dosis no estricta: se permite "Según HDS")
            req = ("cultivo_area","objetivo","producto","fecha","responsable")
            for it in items:
                if all(_has_text(it.get(k)) for k in req):
                    return True
            return False

        # 2) Coerce de lo que venga en el context (si vino textarea como string)
        inv = _as_list_dicts(context.get("inventario_plaguicidas"))
        prog = _as_list_dicts(context.get("programacion_tratamientos"))

        # 3) Fuentes automáticas:
        #    A) bodega/maestro interno (si en tu modelo existen), B) baseline BANANERA
        def _from_bodega_or_master(db, company_id):
            # seguro: si no hay modelos/tablas, retorna []
            try:
                # EJEMPLO: si algún día agregas modelos, mapéalos aquí.
                return []
            except Exception:
                return []

        def _bananera_inv_baseline(ctx):
            # Inventario mínimo de referencia (sin marcas/dosis; seguro para validación)
            # *Campos exigidos:* producto, ingrediente_activo, formulacion, registro, categoria_tox, hds
            # hds: usa la primera ruta/enlace disponible de hds_links/defaults
            hds_links = ctx.get("hds_links") or []
            hds_first = ""
            if isinstance(hds_links, list) and hds_links:
                hds_first = str(hds_links[0])
            elif isinstance(hds_links, str) and hds_links.strip():
                hds_first = hds_links.strip()

            return [
                {
                    "producto": "Fungicida multisitio (Sigatoka)",
                    "ingrediente_activo": "Ej.: Mancozeb u otro multisitio",
                    "formulacion": "WP/SC",
                    "registro": "Según etiqueta/HDS",
                    "categoria_tox": "Según HDS",
                    "hds": hds_first or "HDS física/digital",
                    "proveedor": "Proveedor habitual",
                    "vencimiento": ""
                },
                {
                    "producto": "Fungicida triazol (rotación Sigatoka)",
                    "ingrediente_activo": "Ej.: Propiconazol/Triazol",
                    "formulacion": "EC/SC",
                    "registro": "Según etiqueta/HDS",
                    "categoria_tox": "Según HDS",
                    "hds": hds_first or "HDS física/digital",
                    "proveedor": "Proveedor habitual",
                    "vencimiento": ""
                },
                {
                    "producto": "Insecticida para trips/raspadores",
                    "ingrediente_activo": "Ej.: Spinosad/Imidacloprid u otro autorizado",
                    "formulacion": "SC/SL",
                    "registro": "Según etiqueta/HDS",
                    "categoria_tox": "Según HDS",
                    "hds": hds_first or "HDS física/digital",
                    "proveedor": "Proveedor habitual",
                    "vencimiento": ""
                }
            ]

        def _bananera_prog_baseline(inv_items, ctx):
            # Programación base (sin dosificación): objetivos típicos BANANERA
            resp = ctx.get("agronomo_responsable") or ctx.get("responsable_sst") or "Responsable designado"
            fecha = str(ctx.get("fecha_emision") or "")
            # Elegimos el nombre de producto desde inventario (si hay), caso contrario genéricos
            p1 = (inv_items[0]["producto"] if inv_items else "Fungicida multisitio (Sigatoka)")
            p2 = (inv_items[1]["producto"] if len(inv_items) > 1 else "Fungicida triazol (rotación Sigatoka)")
            p3 = (inv_items[2]["producto"] if len(inv_items) > 2 else "Insecticida para trips/raspadores")
            return [
                {
                    "cultivo_area": "Banano – Lotes en producción",
                    "objetivo": "Manejo preventivo de Sigatoka",
                    "producto": p1,
                    "dosis": "Según etiqueta/HDS",
                    "fecha": fecha,
                    "responsable": resp
                },
                {
                    "cultivo_area": "Banano – Lotes en producción",
                    "objetivo": "Rotación fungicida (Sigatoka)",
                    "producto": p2,
                    "dosis": "Según etiqueta/HDS",
                    "fecha": fecha,
                    "responsable": resp
                },
                {
                    "cultivo_area": "Banano – Áreas con presión de trips",
                    "objetivo": "Control de trips/raspadores",
                    "producto": p3,
                    "dosis": "Según etiqueta/HDS",
                    "fecha": fecha,
                    "responsable": resp
                }
            ]

        # 4) Si lo que vino no es válido, autogenerar
        if not _valid_inv(inv):
            inv = _from_bodega_or_master(db, company_id) or _bananera_inv_baseline(context)
            context["inventario_plaguicidas"] = inv

        if not _valid_prog(prog):
            context["programacion_tratamientos"] = _bananera_prog_baseline(inv, context)

    
    # (opcional, pero útil) si es AGRO-PLAG-01 y falta agrónomo, completar con conocidos
    if str(template_code).strip().upper() == "AGRO-PLAG-01":
        try:
            known = collect_known_fields(db, company_id)
        except Exception:
            known = {}
        if not str(context.get("agronomo_responsable") or "").strip():
            context["agronomo_responsable"] = (
                known.get("agronomo_responsable")
                or context.get("responsable_sst")
                or known.get("responsable_sst")
                or ""
            )
    
    # --- Validación específica CAP-01 (puntos 6 y 7) ---
    if str(template_code).strip().upper() == "CAP-01":
        import json as _json, re as _re
        raw = context.get("presupuesto")
        def _has_budget(p):
            if isinstance(p, list) and any(isinstance(x, dict) for x in p):
                return True
            if isinstance(p, str):
                s = p.strip()
                try:
                    j = _json.loads(s)
                    if isinstance(j, list) and any(isinstance(x, dict) for x in j):
                        context["presupuesto"] = j
                        return True
                except Exception:
                    pass
                m = _re.search(r"[-\d.,]+", s or "")
                if m:
                    return True
            if str(context.get("presupuesto_txt") or "").strip() and str(context.get("presupuesto_total_txt") or "").strip():
                return True
            return False

        if not _has_budget(raw):
            raise HTTPException(status_code=409, detail="EMPTY_REQUIRED_DATA_CAP01: falta presupuesto")
        
    
    
    # --- Validación específica AGRO-PLAG-01 (obligatorios normativos) ---
    elif code_up == "AGRO-PLAG-01":
        missing = []

        # 1) Profesional responsable (Ing. Agrónomo/a)
        if not str(context.get("agronomo_responsable") or "").strip():
            missing.append("agronomo_responsable")

        # 2) HDS accesibles — al menos una ubicación / link
        hds_links = context.get("hds_links")
        has_hds = False
        if isinstance(hds_links, list):
            has_hds = any(str(x or "").strip() for x in hds_links)
        elif isinstance(hds_links, str):
            has_hds = bool(hds_links.strip())
        if not has_hds:
            missing.append("hds_links")

        # 3) Inventario con ≥1 ítem completo
        inv = context.get("inventario_plaguicidas")
        def _ok_item(it):
            req = ["producto","ingrediente_activo","formulacion","registro","categoria_tox","hds"]
            return isinstance(it, dict) and all(str(it.get(k,"")).strip() for k in req)
        inv_ok = isinstance(inv, list) and any(_ok_item(i) for i in inv or [])
        if not inv_ok:
            missing.append("inventario_plaguicidas(≥1 item con producto, ingrediente_activo, formulacion, registro, categoria_tox, hds)")

        if missing:
            raise HTTPException(status_code=409, detail=f"EMPTY_REQUIRED_DATA_AGRO_PLAG01: faltan {', '.join(missing)}")


    context["legal"] = legal_txt
    if legal_txt:
        context["legal_inline"] = legal_txt

           
    # --- NUEVO: sólo para IPERC-01, traer filas IPERC guardadas
    if code_up == "IPERC-01":
        sheet_sel = (data.get("sheet") or "").strip().upper() or None
        q = db.query(IPERCItem).filter(IPERCItem.company_id == company_id)
        if sheet_sel:
            q = q.filter(IPERCItem.sheet == sheet_sel)
        rows_db = q.order_by(IPERCItem.id.asc()).all()
        rows = [_row_to_dict(r) for r in rows_db]

        if not rows:
            raise HTTPException(status_code=409, detail="EMPTY_REQUIRED_DATA")

        def cnt(level): return sum(1 for r in rows if r.get("nr_level") == level)
        def cntnp(tag): return sum(1 for r in rows if r.get("np_level") == tag)
        context.update({
            "rows": rows,
            "total_rows": len(rows),
            "nr_I": cnt("I"), "nr_II": cnt("II"), "nr_III": cnt("III"), "nr_IV": cnt("IV"),
            "np_MA": cntnp("MA"), "np_A": cntnp("A"), "np_M": cntnp("M"), "np_B": cntnp("B"),
            "total_permiso": sum(1 for r in rows if r.get("requires_work_permit")),
            "total_vsalud": sum(1 for r in rows if r.get("needs_health_surveillance")),
            "total_mamb": sum(1 for r in rows if r.get("needs_env_monitoring")),
        })
        
    # --- 🔹 NUEVO BLOQUE PARA PPRL-01 (usa también la matriz IPERC)
    elif code_up == "PPRL-01":
        q = db.query(IPERCItem).filter(IPERCItem.company_id == company_id)
        rows_db = q.order_by(IPERCItem.id.asc()).all()
        rows = [_row_to_dict(r) for r in rows_db]

        # 🔹 Enriquecer con el catálogo IPERC para que tenga controles, EPP, etc.
        try:
            rows = _enrich_iperc_rows_with_catalog(activity, rows)
        except Exception as e:
            log.warning("No se pudo enriquecer IPERC para PPRL-01: %s", e)

        context["rows"] = rows or []

        # ====== PPRL-01: generar textos planos para la plantilla DOCX (Forma C) ======
        import unicodedata as _ud

        def _strip_acc(s: str) -> str:
            s = str(s or "")
            return "".join(ch for ch in _ud.normalize("NFD", s) if _ud.category(ch) != "Mn")

        def _norm(s: str) -> str:
            return _strip_acc(s).upper().strip()

        def _as_list(x):
            if not x: return []
            if isinstance(x, (list, tuple, set)): return list(x)
            return [x]

        def _bullets(items):
            items = [str(i).strip() for i in items if str(i or "").strip()]
            return "\n".join(f"- {i}" for i in items) if items else "- No observado"

        def _controls(r, keys):
            out = []
            for k in keys:
                out += _as_list(r.get(k))
            return out

        def _group(rows, names):
            want = { _norm(n) for n in _as_list(names) }
            return [r for r in rows if _norm(r.get("hazard_group")) in want]

        def _set_group(prefix, names):
            g = _group(rows, names)
            context[f"{prefix}_peligros_txt"] = _bullets([f"{r.get('hazard')} ({r.get('process')} – {r.get('job')})" for r in g])
            context[f"{prefix}_existentes_txt"] = _bullets([c for r in g for c in _controls(r, [
                "controls_existing_engineering","controls_existing_admin","controls_existing_epp"
            ])])
            context[f"{prefix}_prop_txt"] = _bullets([c for r in g for c in _controls(r, [
                "controls_proposed","controls_planned_engineering","controls_planned_admin","controls_planned_epp"
            ])])

        # Grupos (acepta sin/ con acentos)
        _set_group("fisicos", "FISICOS")
        _set_group("quimicos", "QUIMICOS")
        _set_group("biologicos", "BIOLOGICOS")
        _set_group("ergo", "ERGONOMICOS")
        _set_group("psico", "PSICOSOCIALES")
        _set_group("seg", ["MECANICOS", "ELECTRICOS", "SEGURIDAD"])

        # Plan de acción (7 columnas, una línea por acción)
        all_actions = []
        for r in rows:
            props = _controls(r, [
                "controls_proposed","controls_planned_engineering","controls_planned_admin","controls_planned_epp"
            ])
            for c in props:
                all_actions.append((r, c))

        if all_actions:
            # placeholders de texto (compatibilidad con la plantilla actual)
            context["pa_control_txt"]   = "\n".join(c for (r, c) in all_actions)
            context["pa_riesgo_txt"]    = "\n".join(f"{r.get('hazard_group')} – {r.get('hazard')} ({r.get('process')} – {r.get('job')})" for (r, c) in all_actions)
            context["pa_responsable_txt"]= "\n".join([(data.get("responsable_sst") or "Delegado/a de SST") for _ in all_actions])
            context["pa_recursos_txt"]  = "\n".join(["Recursos internos SST / presupuesto del área" for _ in all_actions])
            context["pa_f_inicio_txt"]  = "\n".join([str(data.get("fecha_emision") or "") for _ in all_actions])
            context["pa_f_fin_txt"]     = "\n".join([str(data.get("periodo_vigencia") or "") for _ in all_actions])
            context["pa_evidencia_txt"] = "\n".join(["Registro: procedimiento/capacitación/acta/inspección" for _ in all_actions])

            # NUEVO: lista de filas para repetir en Word (una por acción)
            context["pa_rows"] = [
                {
                    "control": c,
                    "riesgo": f"{r.get('hazard_group')} – {r.get('hazard')} ({r.get('process')} – {r.get('job')})",
                    "responsable": data.get("responsable_sst") or "Delegado/a de SST",
                    "recursos": "Recursos internos SST / presupuesto del área",
                    "f_inicio": str(data.get("fecha_emision") or ""),
                    "f_fin": str(data.get("periodo_vigencia") or ""),
                    "evidencia": "Registro: procedimiento/capacitación/acta/inspección",
                }
                for (r, c) in all_actions
            ]
        else:
            # placeholders de texto (compatibilidad)
            context["pa_control_txt"]   = "No observado"
            context["pa_riesgo_txt"]    = "No aplica"
            context["pa_responsable_txt"]= data.get("responsable_sst") or "Delegado/a de SST"
            context["pa_recursos_txt"]  = "Recursos internos SST"
            context["pa_f_inicio_txt"]  = str(data.get("fecha_emision") or "")
            context["pa_f_fin_txt"]     = str(data.get("periodo_vigencia") or "")
            context["pa_evidencia_txt"] = "Sin registro"

            # NUEVO: una fila “vacía” para que Word muestre algo si no hay acciones
            context["pa_rows"] = [
                {
                    "control": "No observado",
                    "riesgo": "No aplica",
                    "responsable": data.get("responsable_sst") or "Delegado/a de SST",
                    "recursos": "Recursos internos SST",
                    "f_inicio": str(data.get("fecha_emision") or ""),
                    "f_fin": str(data.get("periodo_vigencia") or ""),
                    "evidencia": "Sin registro",
                }
            ]

    elif code_up == "CAP-01":
        # 1) Traer IPERC de BD (mismo patrón que PPRL/IPERC)
        q = db.query(IPERCItem).filter(IPERCItem.company_id == company.id)
        rows_db = q.order_by(IPERCItem.id.asc()).all()
        rows = [_row_to_dict(r) for r in rows_db]

        # 2) Enriquecer con catálogo BANANERA (si falla, no romper)
        try:
            rows = _enrich_iperc_rows_with_catalog(activity, rows)
        except Exception:
            pass

        # 3) Derivar temas de capacitación SÓLO si en controles planificados se detectan
        #    temas/acciones tipo capacitación. Evitamos inventar temas no presentes.
        def _as_list(x):
            if x is None: return []
            if isinstance(x, list): return x
            if isinstance(x, (tuple, set)): return list(x)
            return [x]
        def _controls_planned(r):
            return (
                _as_list(r.get("controls_planned_engineering")) +
                _as_list(r.get("controls_planned_admin")) +
                _as_list(r.get("controls_planned_epp")) +
                _as_list(r.get("controls_proposed"))
            )

        # Palabras clave -> tema estandarizado
        TOPIC_MAP = [
            ("extintor", "Uso de extintores y fuego incipiente"),
            ("simulac", "Plan de emergencias y simulacros"),
            ("epp", "Uso, selección y mantenimiento de EPP"),
            ("plaguicid", "Manejo seguro de plaguicidas y HDS"),
            ("eléctr", "Seguridad eléctrica"),
            ("electr", "Seguridad eléctrica"),
            ("ergono", "Ergonomía y manipulación de cargas"),
            ("psicosocial", "Prevención de riesgos psicosociales y liderazgo"),
            ("estrés térmico", "Prevención de estrés térmico"),
            ("estres termico", "Prevención de estrés térmico"),
            ("fertiliz", "Manejo de fertilizantes"),
            ("bodega", "Gestión segura de bodegas de químicos"),
            ("almac", "Gestión segura de bodegas de químicos"),
        ]

        def _match_topic(txt: str) -> str | None:
            t = (txt or "").lower()
            for kw, tema in TOPIC_MAP:
                if kw in t:
                    return tema
            return None

        # 4) Construir set de temas y referencias IPERC (evita duplicados)
        temas = {}
        for r in rows or []:
            ref_iperc = f"{r.get('process','')} – {r.get('job','')} – {r.get('hazard_group','')}/{r.get('hazard','')}".strip()
            for c in _controls_planned(r):
                tema = _match_topic(str(c))
                if not tema:
                    continue
                if tema not in temas:
                    temas[tema] = {
                        "origen": "PPRL:PlanAccion",
                        "tema": tema,
                        "detalle": "",
                        "referencia_iperc": ref_iperc,
                        "poblacion": "Según puesto/proceso relacionado",
                        "n_participantes": "Objetivo",
                        "periodicidad": "Anual",
                        "duracion_h": 2,
                        "modalidad": "Presencial",
                        "proveedor": "SST interno/externo",
                        "mes": "",  # se asigna abajo,
                        "evidencia": "Lista, evaluación",
                        "indicador": "Cobertura mayor o igual a 90 por ciento",
                        "source_doc": data.get("cap_fte_pprl") or "",
                        "source_section": "Plan de acción",
                    }

        # Si NO se detectó nada en controles, intenta con Sección textual de PPRL si el JSON la trajo
        # (cuando el integrador haya guardado 'pprl_01'->'rows' con 'text_section_7' o similar)
        # ==> mantener simple: si no hay temas, dejamos lista vacía (la tabla mostrará “— Sin datos —”).

        # 5) Asignar meses simples (sin % ni formatos extra)
        mes_orden = [
            "Enero","Febrero","Marzo","Abril","Mayo","Junio",
            "Julio","Agosto","Septiembre","Octubre","Noviembre","Diciembre"
        ]
        cap_list = list(temas.values())
        for i, it in enumerate(cap_list):
            if not it.get("mes"):
                it["mes"] = mes_orden[i % 12]

        # 6) Programación mensual agrupada
        prog = []
        for it in cap_list:
            prog.append({
                "mes": it["mes"],
                "actividades": f"{it['tema']} ({it['origen']})",
                "areas": "Según IPERC",
                "responsable": data.get("responsable_sst") or "",
                "observaciones": "",
            })

        # 7) Devolver al contexto para que el branch CAP-01 en render_document construya las tablas
        context["capacitaciones"] = cap_list
        context["programacion_mensual"] = prog
    
    elif str(template_code).upper() == "EPP-01":
        # 1) Traer IPERC de BD y enriquecer con catálogo (no romper si falla)
        q = db.query(IPERCItem).filter(IPERCItem.company_id == company.id)
        rows_db = q.order_by(IPERCItem.id.asc()).all()
        rows = [_row_to_dict(r) for r in rows_db]
        try:
            rows = _enrich_iperc_rows_with_catalog(activity, rows)
        except Exception:
            pass

        # 2) Agrupar por (process, job) y consolidar peligros/EPP
        def _as_list(x):
            if x is None: return []
            if isinstance(x, list): return x
            if isinstance(x, (tuple, set)): return list(x)
            return [x]

        buckets = {}  # (process, job) -> {"hazards": set(), "epp": set()}
        for r in rows or []:
            proc = str(r.get("process","") or "").strip()
            job  = str(r.get("job","") or "").strip()
            key = (proc, job)
            b = buckets.setdefault(key, {"hazards": set(), "epp": set()})
            hz = "/".join(filter(None, [str(r.get("hazard_group","") or ""), str(r.get("hazard","") or "")])).strip("/")
            if hz:
                b["hazards"].add(hz)

            # Prioriza EPP crítico, si no, existentes y/o planificados
            epp_src = (_as_list(r.get("critical_epp"))
                       or _as_list(r.get("controls_existing_epp"))
                       or _as_list(r.get("controls_planned_epp")))
            for e in epp_src:
                if str(e).strip():
                    b["epp"].add(str(e).strip())

        # 3) Construir matriz EPP
        epp_matrix = []
        for (proc, job), b in sorted(buckets.items(), key=lambda x: (x[0][0], x[0][1])):
            epp_matrix.append({
                "process": proc,
                "job": job,
                "hazards_txt": ", ".join(sorted(b["hazards"])) if b["hazards"] else "",
                "epp_txt":     ", ".join(sorted(b["epp"]))     if b["epp"]     else "",
                "reposicion": "Por desgaste, daño, pérdida de eficacia o cambio de tarea",
                "responsable_entrega": (data.get("responsable_sst") or "Bodega/SST"),
                "evidencia": "Registro de entrega EPP-01"
            })

        # 4) Registros de entrega (formato de firma) — un renglón por puesto
        epp_registros = []
        for row in epp_matrix:
            epp_registros.append({
                "nombre": "",
                "puesto": row["job"],
                "epp_entregado": row["epp_txt"],
                "talla": "",
                "cantidad": "",
                "fecha": "",
                "firma_trabajador": "",
                "firma_responsable": data.get("responsable_sst") or ""
            })

        # 5) Contexto para el render DOCX
        context["epp_matrix"] = epp_matrix
        context["epp_registros"] = epp_registros
    
        
    elif code_up == "EMG-01":
        # 1) Validación mínima (solo datos que NO provienen de otros documentos)
        required = []
        if not str(data.get("numero_emergencia_interno") or "").strip():
            required.append("numero_emergencia_interno")
        if not str(data.get("contacto_bomberos") or "").strip():
            required.append("contacto_bomberos")
        if required:
            raise HTTPException(status_code=409, detail=f"EMPTY_REQUIRED_DATA_EMG01: faltan {', '.join(required)}")

        # 2) Traer IPERC y enriquecer (igual que CAP-01/PPRL-01)
        q = db.query(IPERCItem).filter(IPERCItem.company_id == company.id)
        rows_db = q.order_by(IPERCItem.id.asc()).all()
        rows = [_row_to_dict(r) for r in rows_db]
        try:
            rows = _enrich_iperc_rows_with_catalog(activity, rows)
        except Exception:
            pass

        # 3) Derivar escenarios de emergencia desde IPERC/PPRL (clasificación sencilla por palabras clave)
        def _as_list(x):
            if x is None: return []
            if isinstance(x, list): return x
            if isinstance(x, (tuple, set)): return list(x)
            return [x]

        KEYMAP = [
            ("INCENDIO/FUEGO", ("incend", "fuego", "combust", "chisp", "explos")),
            ("DERRAME/QUÍMICOS", ("derram", "quimic", "plaguicid", "toxic", "hds")),
            ("ELÉCTRICO", ("electr", "tablero", "energiz", "shock", "corto")),
            ("ESTRUCTURAL/MECÁNICO", ("mecan", "caida objetos", "volcad", "atrap", "golpe")),
            ("NATURAL", ("sismo", "inund", "vendav", "tormenta", "desliz")),
        ]

        def _classify(text: str) -> str:
            t = (text or "").lower()
            for clase, kws in KEYMAP:
                if any(k in t for k in kws):
                    return clase
            return "OTROS"

        escenarios = {}
        for r in rows or []:
            ref = f"{r.get('process','')} – {r.get('job','')} – {r.get('hazard_group','')}/{r.get('hazard','')}".strip()
            hz = " ".join([str(r.get("hazard_group","")), str(r.get("hazard",""))])
            clase = _classify(hz)
            # seleccionar controles existentes + EPP crítico como equipos
            controles = _as_list(r.get("controls_existing_engineering")) + _as_list(r.get("controls_existing_admin"))
            equipos = _as_list(r.get("critical_epp")) or _as_list(r.get("controls_existing_epp"))
            key = (clase, ref)
            if key not in escenarios:
                escenarios[key] = {
                    "escenario": r.get("hazard") or "Evento",
                    "clasificacion": clase,
                    "referencia_iperc": ref,
                    "controles_existentes": ", ".join(str(x) for x in controles[:3]),
                    "equipos_requeridos": ", ".join(str(x) for x in equipos[:3]),
                    "roles": data.get("coordinador_emergencias") or data.get("responsable_sst") or "",
                    "punto_encuentro": data.get("puntos_encuentro") or "",
                    "evidencia": "Acta de simulacro / registro de evento"
                }

        escenarios_list = list(escenarios.values())

        # 4) Programación de simulacros: tomar 2-3 escenarios más frecuentes
        meses = ["Febrero","Junio","Noviembre"]
        sim_rows = []
        for i, esc in enumerate(escenarios_list[:3]):
            sim_rows.append({
                "mes": meses[i % len(meses)],
                "escenario": f"{esc['clasificacion']} – {esc['escenario']}",
                "objetivo": "Verificar tiempos de respuesta y roles",
                "responsable": data.get("coordinador_emergencias") or data.get("responsable_sst") or "",
                "observaciones": ""
            })

        # 5) Contactos externos (tabla)
        contactos = []
        if str(data.get("contacto_bomberos") or "").strip():
            contactos.append({"entidad":"Bomberos", "contacto": data.get("contacto_bomberos"), "telefono":"", "obs":""})
        if str(data.get("contacto_salud") or "").strip():
            contactos.append({"entidad":"Salud", "contacto": data.get("contacto_salud"), "telefono":"", "obs":""})
        if str(data.get("contacto_policia") or "").strip():
            contactos.append({"entidad":"Policía", "contacto": data.get("contacto_policia"), "telefono":"", "obs":""})
        if str(data.get("contacto_municipio") or "").strip():
            contactos.append({"entidad":"Municipio", "contacto": data.get("contacto_municipio"), "telefono":"", "obs":""})

        # 6) Entregar al contexto para que render_document cree subdocs
        context["emg_escenarios"] = escenarios_list
        context["emg_simulacros"] = sim_rows
        context["emg_contactos"] = contactos
        
    
    elif code_up == "PSICO-01":
        # 1) Trae campos conocidos (…)
        try:
            known = collect_known_fields(db, company_id) or {}
        except Exception:
            known = {}
        for k in ["razon_social","ruc","representante_legal","responsable_sst","actividad","riesgo","trabajadores"]:
            if not str(context.get(k) or "").strip():
                v = known.get(k) or getattr(company, "name" if k=="razon_social" else k, "")
                if v: context[k] = v
                
            # === NUEVO: preparar marcas de cuestionario y sumatorias exactamente como la hoja ===
        try:
            _prepare_psico_marks(context)
        except Exception:
            # no romper si el front aún no envía todo
            pass

        # 1.1) Cálculo psicosocial desde psico_respuestas + defaults umbrales/acciones
        try:
            psico_calc = _compute_psico(context)
        except Exception:
            psico_calc = {}
        for k, v in (psico_calc or {}).items():
            if v not in (None, [], {}) and k not in context:
                context[k] = v

        # 2) Defaults técnicos (solo si faltan)
        context.setdefault("anio", datetime.utcnow().year)
        context.setdefault("nota_confidencialidad",
            "La aplicación del cuestionario y tratamiento de resultados se realizan de forma confidencial, "
            "limitada a fines preventivos, sin uso disciplinario ni individualización de respuestas.")
        context.setdefault("instrumento", "Cuestionario de evaluación de riesgo psicosocial validado por el órgano rector")
        context.setdefault("cobertura", "Toda la población trabajadora, según categorías ocupacionales")

        # 3) Equipo (… como lo tienes …)

        # 4) Cronograma (… como lo tienes …)

        # 5) Resultados sintéticos SOLO si aún no hay (usar cálculo si existe)
        if not context.get("psico_resultados"):
            if context.get("psico_dimensiones"):
                context["psico_resultados"] = [
                    {"factor": d["factor"], "nivel": d["nivel"], "hallazgos": "", "poblacion": ""}
                    for d in context["psico_dimensiones"]
                ]
            else:
                context["psico_resultados"] = [
                    {"factor":"Demandas de trabajo","nivel":"Medio","hallazgos":"Ritmo y carga en picos estacionales","poblacion":"Operativa"},
                    {"factor":"Control sobre el trabajo","nivel":"Bajo","hallazgos":"Autonomía adecuada en la mayoría de puestos","poblacion":"General"},
                ]

        # 6) Plan de intervención SOLO si aún no hay (usar plan sugerido si existe)
        if not context.get("psico_plan"):
            if context.get("psico_plan_sugerido"):
                context["psico_plan"] = context["psico_plan_sugerido"]
            else:
                fi = str(context.get("fecha_emision") or "")
                context["psico_plan"] = [
                    {"medida":"Adecuar cargas y pausas activas","factor":"Demandas de trabajo","responsable":"Jefaturas / SST","recursos":"Recursos internos","f_inicio":fi,"f_fin":"", "evidencia":"Registro de pausas"},
                    {"medida":"Capacitación en liderazgo y comunicación","factor":"Relaciones","responsable":"Talento Humano","recursos":"Proveedor externo","f_inicio":fi,"f_fin":"", "evidencia":"Listas y evaluación"},
                ]
        # 7) Indicadores
        if not context.get("psico_indicadores"):
            context["psico_indicadores"] = [
                {"indicador":"Cobertura de cuestionario","formula":"(Nº encuestas / Nº trabajadores) x 100","meta":"≥ 80%","fuente":"Listas de aplicación","frecuencia":"Anual"},
                {"indicador":"Ejecución del plan de intervención","formula":"(Actividades ejecutadas / planificadas) x 100","meta":"≥ 90%","fuente":"Cronograma / registros","frecuencia":"Trimestral"},
                {"indicador":"Eficacia percibida","formula":"% de mejora reportada","meta":"Tendencia al alza","fuente":"Encuesta breve post-intervención","frecuencia":"Anual"},
            ]

    
    elif str(template_code).upper() == "AGRO-PLAG-01":
        # Los campos ya se autocompletaron arriba (línea ~1342) desde docs anteriores
        
        # 1) Traer IPERC y filtrar riesgos QUÍMICOS (incluye plaguicidas)
        q = db.query(IPERCItem).filter(IPERCItem.company_id == company.id)
        rows_db = q.order_by(IPERCItem.id.asc()).all()
        rows = [_row_to_dict(r) for r in rows_db]
        try:
            rows = _enrich_iperc_rows_with_catalog(activity, rows)
        except Exception:
            pass

        chem_rows = [r for r in (rows or []) if str(r.get("hazard_group","")).strip().upper().startswith(("QUIM", "QUÍM"))]

        def _as_list(x):
            if x is None: return []
            if isinstance(x, list): return x
            if isinstance(x, (tuple, set)): return list(x)
            return [x]

        # --- Normalizador de EPP canónico (evita duplicados y genéricos ambiguos)
        def _canon_epp(name: str) -> str:
            s = (name or "").lower()
            if any(k in s for k in ("respir", "mascar", "filtro")):
                return "Respirador con filtro adecuado al producto (según HDS)"
            if "guante" in s:
                return "Guantes químicos compatibles (según HDS)"
            if any(k in s for k in ("gafa","pantalla","facial","ocular")):
                return "Protección ocular/facial (según HDS)"
            if "bota" in s:
                return "Botas impermeables"
            if any(k in s for k in ("overol","mandil","delantal","ropa")):
                return "Ropa de protección química (según HDS)"
            return name.strip()

        # 2) EPP obligatorio (prioriza crítico; si no, existentes/planned) + normalización
        epp_set = set()
        for r in chem_rows:
            src = (_as_list(r.get("critical_epp"))
                or _as_list(r.get("controls_existing_epp"))
                or _as_list(r.get("controls_planned_epp")))
            for e in src:
                s = _canon_epp(str(e).strip())
                if s: epp_set.add(s)
        if epp_set and not context.get("epp_obligatorio"):
            context["epp_obligatorio"] = sorted(epp_set)

        # 3) Temas de capacitación derivados (con defaults)
        TOPICS = [
            ("plaguicid", "Manejo seguro de plaguicidas y HDS"),
            ("hds", "Interpretación de HDS y comunicación de peligros (GHS)"),
            ("almac", "Almacenamiento seguro de agroquímicos"),
            ("mezcla", "Preparación y dosificación segura"),
            ("calibr", "Calibración de equipos de aplicación"),
            ("derrames", "Respuesta a derrames y contención"),
            ("epp", "Uso, mantenimiento y reposición de EPP específico"),
            ("triple lavado", "Triple lavado y gestión de envases"),
            ("expos", "Vías de exposición y medidas preventivas"),
        ]
        def _match_any(text):
            t = (text or "").lower()
            for kw, tema in TOPICS:
                if kw in t:
                    return tema
            return None

        temas = set()
        for r in chem_rows:
            for c in (_as_list(r.get("controls_planned_admin"))
                    + _as_list(r.get("controls_planned_engineering"))
                    + _as_list(r.get("controls_planned_epp"))
                    + _as_list(r.get("controls_proposed"))):
                t = _match_any(str(c))
                if t: temas.add(t)

        if not temas:
            temas = {
                "Manejo seguro de plaguicidas y HDS",
                "Interpretación de HDS y comunicación de peligros (GHS)",
                "Almacenamiento seguro de agroquímicos",
                "Preparación y dosificación segura",
                "Calibración de equipos de aplicación",
                "Respuesta a derrames y contención",
                "Uso, mantenimiento y reposición de EPP específico",
                "Triple lavado y gestión de envases",
            }
        if not context.get("capacitacion_temas"):
            context["capacitacion_temas"] = sorted(temas)

        # 4) Vigilancia de la salud (sugerencias + criterio de acción)
        if not context.get("vigilancia_items"):
            context["vigilancia_items"] = [
                "Ficha clínica y anamnesis de exposición a plaguicidas",
                "Examen preocupacional y periódico con énfasis en exposición a químicos",
                "Biomarcadores según ingrediente activo (p. ej., colinesterasa con línea base y comparativa)",
                "Criterio de acción: variación ≥ 20% vs basal → retiro temporal y evaluación",
                "Evaluación de aptitud y restricciones por puesto",
            ]

        # 5) Procedimientos escritos mínimos (REI + bitácoras + calibración)
        if not context.get("procedimientos_escritos"):
            context["procedimientos_escritos"] = [
                "Recepción, almacenamiento y control de inventario de agroquímicos",
                "Preparación de mezcla, dosificación y aplicación",
                "Señalización de áreas tratadas y control del Periodo de Reingreso (REI)",
                "Limpieza de equipos, triple lavado y disposición de envases",
                "Bitácoras: recepción, inventario, mezclas, calibraciones y aplicaciones",
                "Respuesta a emergencias: derrames, intoxicaciones y fuego",
                "Gestión de HDS y etiquetado GHS",
            ]

        # 6) Lineamientos de almacenamiento (refuerzo normativo)
        if not context.get("almacenamiento_lineamientos"):
            context["almacenamiento_lineamientos"] = [
                "Bodega exclusiva, ventilada, con dique de contención y piso impermeable",
                "Segregación por compatibilidad; prohibido almacenar con alimentos o EPP limpios",
                "Señalización GHS, control de acceso y registro de entradas/salidas",
                "Estanterías resistentes y bandejas de contención",
                "Plan de emergencias con extintores adecuados y kit de derrames",
            ]

        # 8) Indicadores sugeridos (REI + calibración + HDS)
        if not context.get("indicadores"):
            context["indicadores"] = [
                "Cobertura de capacitación en manejo de plaguicidas ≥ 90%",
                "HDS disponibles y accesibles al 100% de productos",
                "Cumplimiento de REI (señalización y control) ≥ 95%",
                "Calibraciones registradas antes de cada campaña ≥ 95%",
                "Triple lavado y disposición de envases conforme ≥ 95%",
                "Uso de EPP según HDS ≥ 95%",
                "Incidentes por exposición a plaguicidas = 0",
            ]

        # 8) AUTOCOMPLETAR campos básicos de empresa SOLO si no vienen del usuario
        # Esto se hace AL FINAL para no interferir con la derivación inteligente del IPERC
        basic_fields = {
            "razon_social": getattr(company, "name", ""),
            "ruc": getattr(company, "ruc", ""),
        }
        
        # Intentar obtener otros campos de documentos anteriores
        try:
            known = collect_known_fields(db, company_id)
            for field in ["direccion", "representante_legal", "responsable_sst", 
                          "medico_ocupacional", "responsable_bodega"]:
                if field in known:
                    basic_fields[field] = known[field]
        except Exception:
            pass
        
        # Aplicar campos básicos SOLO si no están en context o están vacíos
        for k, v in basic_fields.items():
            if not context.get(k) or context.get(k) in ("", [], {}):
                context[k] = v
   
        # Alias para plantillas antiguas (AGRO-PLAG-01 usa {{ company }})
        context.setdefault("company", context.get("razon_social") or getattr(company, "name", ""))
                        
    # 2) Secuencia/versión y 3) código con versión
    ver = next_version(db, company_id=company_id, template_code=template_code)
    code_full = compose_code_versioned(template_code, datetime.utcnow(), ver)

    # 4) Renderizar contenido (DOCX si existe, si no JSON)
    content, mime, ext = render_document(tpl, context)

    # 1) Normaliza extensión
    ext_final = (ext or "docx").lstrip(".").lower()

    # 2) Si dice ser DOCX, valida que realmente sea ZIP (PK…)
    if ext_final == "docx":
        is_zip = False
        try:
            if isinstance(content, (bytes, bytearray)) and content[:2] == b"PK":
                zipfile.ZipFile(io.BytesIO(content)).testzip()
                is_zip = True
        except Exception:
            is_zip = False

        if not is_zip:
            # No es un .docx válido → degrada a JSON (contenido legible)
            ext_final = "json"
            mime = "application/json"
            if isinstance(content, (bytes, bytearray)):
                content = json.dumps(
                    {"error": "docx-invalid", "context": {"title": tpl.title, "code": tpl.code}},
                    ensure_ascii=False
                ).encode("utf-8")
            elif isinstance(content, dict):
                content = json.dumps(content, ensure_ascii=False, indent=2).encode("utf-8")
            else:
                content = json.dumps({"data": str(content)}, ensure_ascii=False).encode("utf-8")

    # 2 bis) Reparar DOCX para máxima compatibilidad con Word (solo PPRL-01)
    if ext_final == "docx" and str(template_code).upper() == "PPRL-01":
        try:
            from docx import Document as _WordDoc
            # content ya debería ser bytes; lo envolvemos por seguridad
            buf = io.BytesIO(
                content if isinstance(content, (bytes, bytearray)) else bytes(content)
            )
            d = _WordDoc(buf)
            out2 = io.BytesIO()
            d.save(out2)
            content = out2.getvalue()
        except Exception:
            # Si algo falla aquí, dejamos el contenido original
            pass

    # 🔵 2.5) Normaliza el MIME por extensión
    if ext_final == "xlsx":
        mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    elif ext_final == "docx":
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif ext_final == "pdf":
        mime = "application/pdf"
    elif ext_final == "json":
        mime = "application/json"
    else:
        mime = mime or "application/octet-stream"


    # Asegura bytes si aún no lo son
    if not isinstance(content, (bytes, bytearray)):
        content = str(content).encode("utf-8")

    filename = f"{code_full}.{ext_final}"
    path = save_bytes(company_id=company_id, filename=filename, content=content, user_id=user_id, doc_type="DOCS")

    meta = {
        "form": data,
        "activity": activity,
        "risk": risk_level,
        "workers": workers,
        **extra_meta,
    }

    doc = Document(
        company_id=company_id,
        kind=template_code,
        title=code_full,
        storage_path=path,
        mime=mime,
        requirement_code=template_code,
        period_year=datetime.utcnow().year,
        seq=ver,
        meta=meta,
    )

    db.add(doc)
    db.commit()
    db.refresh(doc)
    return doc
